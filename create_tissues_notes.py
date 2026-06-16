#!/usr/bin/env python3
"""
Generate comprehensive study notes on "The Tissues" for Human Anatomy and Physiology.
Creates a visually attractive DOCX file suitable for exam preparation.
Sources: Guyton, Ross & Wilson, Tortora, Sembulingam, Junqueira's Basic Histology
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import copy


def set_cell_shading(cell, color):
    """Set background shading for a table cell."""
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)


def set_cell_border(cell, **kwargs):
    """Set cell border properties."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}>'
                          f'<w:top w:val="single" w:sz="4" w:color="000000"/>'
                          f'<w:bottom w:val="single" w:sz="4" w:color="000000"/>'
                          f'<w:left w:val="single" w:sz="4" w:color="000000"/>'
                          f'<w:right w:val="single" w:sz="4" w:color="000000"/>'
                          f'</w:tcBorders>')
    tcPr.append(tcBorders)


def add_table_borders(table):
    """Add borders to all cells in a table."""
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>') 
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'<w:top w:val="single" w:sz="6" w:color="2E4057"/>'
        f'<w:left w:val="single" w:sz="6" w:color="2E4057"/>'
        f'<w:bottom w:val="single" w:sz="6" w:color="2E4057"/>'
        f'<w:right w:val="single" w:sz="6" w:color="2E4057"/>'
        f'<w:insideH w:val="single" w:sz="4" w:color="2E4057"/>'
        f'<w:insideV w:val="single" w:sz="4" w:color="2E4057"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)


def create_styled_table(doc, headers, rows, header_color="1B4F72"):
    """Create a styled table with header shading."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_table_borders(table)

    # Header row
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(header)
        run.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.font.size = Pt(10)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, header_color)

    # Data rows
    for row_idx, row_data in enumerate(rows):
        for col_idx, cell_text in enumerate(row_data):
            cell = table.rows[row_idx + 1].cells[col_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(cell_text))
            run.font.size = Pt(9)
            if row_idx % 2 == 0:
                set_cell_shading(cell, "EBF5FB")

    return table


def add_highlight_box(doc, title, content, color="FEF9E7", border_color="F39C12"):
    """Add a highlighted box for important points."""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.rows[0].cells[0]
    set_cell_shading(cell, color)

    # Add border
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'<w:top w:val="single" w:sz="12" w:color="{border_color}"/>'
        f'<w:bottom w:val="single" w:sz="12" w:color="{border_color}"/>'
        f'<w:left w:val="single" w:sz="12" w:color="{border_color}"/>'
        f'<w:right w:val="single" w:sz="12" w:color="{border_color}"/>'
        f'</w:tcBorders>'
    )
    tcPr.append(tcBorders)

    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(f"  {title}")
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x8B, 0x45, 0x13)

    if isinstance(content, list):
        for item in content:
            p = cell.add_paragraph()
            run = p.add_run(f"  {item}")
            run.font.size = Pt(10)
    else:
        p = cell.add_paragraph()
        run = p.add_run(f"  {content}")
        run.font.size = Pt(10)

    doc.add_paragraph()


def add_section_divider(doc, color="1B4F72"):
    """Add a visual section divider."""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.rows[0].cells[0]
    cell.text = ""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'<w:bottom w:val="single" w:sz="18" w:color="{color}"/>'
        f'</w:tcBorders>'
    )
    tcPr.append(tcBorders)
    doc.add_paragraph()


def add_heading_styled(doc, text, level=1):
    """Add a styled heading with custom colors."""
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        if level == 1:
            run.font.color.rgb = RGBColor(0x1B, 0x4F, 0x72)
            run.font.size = Pt(22)
        elif level == 2:
            run.font.color.rgb = RGBColor(0x15, 0x4F, 0x0B)
            run.font.size = Pt(16)
        elif level == 3:
            run.font.color.rgb = RGBColor(0x6C, 0x3A, 0x83)
            run.font.size = Pt(13)
    return heading


def add_bullet_points(doc, items, bold_prefix=False):
    """Add formatted bullet points."""
    for item in items:
        p = doc.add_paragraph(style='List Bullet')
        if bold_prefix and ':' in item:
            parts = item.split(':', 1)
            run = p.add_run(parts[0] + ':')
            run.bold = True
            run.font.size = Pt(10)
            run = p.add_run(parts[1])
            run.font.size = Pt(10)
        else:
            run = p.add_run(item)
            run.font.size = Pt(10)


def add_numbered_list(doc, items):
    """Add formatted numbered list."""
    for i, item in enumerate(items, 1):
        p = doc.add_paragraph(style='List Number')
        run = p.add_run(item)
        run.font.size = Pt(10)


def add_body_text(doc, text, bold=False, italic=False):
    """Add body text with formatting."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    run.bold = bold
    run.italic = italic
    return p


def create_title_page(doc):
    """Create an attractive title page."""
    # Add spacing
    for _ in range(3):
        doc.add_paragraph()

    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("THE TISSUES")
    run.bold = True
    run.font.size = Pt(36)
    run.font.color.rgb = RGBColor(0x1B, 0x4F, 0x72)

    # Subtitle
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Human Anatomy & Physiology")
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(0x2E, 0x86, 0xC1)
    run.italic = True

    doc.add_paragraph()

    # Info box
    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = info.add_run("Comprehensive Study Notes for Exam Preparation")
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x6C, 0x3A, 0x83)

    doc.add_paragraph()

    sources = doc.add_paragraph()
    sources.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sources.add_run("Sources: Guyton & Hall | Tortora | Ross & Wilson | Sembulingam | Junqueira")
    run.font.size = Pt(11)
    run.italic = True
    run.font.color.rgb = RGBColor(0x56, 0x6D, 0x7E)

    doc.add_paragraph()

    exams = doc.add_paragraph()
    exams.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = exams.add_run("For: NEET | AIIMS | JIPMER | GPAT | USMLE | University Exams")
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x78, 0x28, 0x1F)
    run.bold = True

    doc.add_page_break()


def add_introduction(doc):
    """Add introduction section."""
    add_heading_styled(doc, "INTRODUCTION TO TISSUES", 1)
    add_section_divider(doc)

    add_body_text(doc, "Histology (Greek: histos = tissue, logos = study) is the study of tissues. "
                  "A tissue is defined as a group of similar cells along with their intercellular "
                  "substance that perform a specific function. The study of tissues forms the "
                  "foundation of understanding organ structure and function.")

    doc.add_paragraph()
    add_heading_styled(doc, "Definition & Importance", 2)
    add_bullet_points(doc, [
        "Tissue (Latin: texere = to weave): A group of cells similar in structure and function",
        "Tissues bridge the gap between cells and organs in the structural hierarchy",
        "Rudolf Virchow (1858): 'Omnis cellula e cellula' - every cell from a cell",
        "Marie Francois Xavier Bichat (Father of Histology): First classified tissues without microscope",
        "Structural hierarchy: Chemical level > Cellular level > Tissue level > Organ level > System level > Organism"
    ], bold_prefix=True)

    doc.add_paragraph()
    add_heading_styled(doc, "Classification of Tissues", 2)
    add_body_text(doc, "Based on structure and function, tissues are classified into FOUR primary types:")

    create_styled_table(doc,
        ["Type", "Origin", "Primary Function", "Key Feature"],
        [
            ["Epithelial", "All 3 germ layers", "Covering & Lining", "Minimal intercellular matrix"],
            ["Connective", "Mesoderm (mesenchyme)", "Support & Binding", "Abundant intercellular matrix"],
            ["Muscular", "Mesoderm", "Movement & Contraction", "Contractile proteins"],
            ["Nervous", "Ectoderm (neuroectoderm)", "Communication & Control", "Electrical excitability"],
        ])

    doc.add_paragraph()
    add_highlight_box(doc, "MNEMONIC: 'ECMN' - Every Cell Must Navigate",
                     "Epithelial, Connective, Muscular, Nervous - the four primary tissue types")

    doc.add_page_break()


def add_epithelial_tissue(doc):
    """Add comprehensive epithelial tissue section."""
    add_heading_styled(doc, "SECTION 1: EPITHELIAL TISSUE", 1)
    add_section_divider(doc)

    add_body_text(doc, "Epithelial tissue (epithelium) covers body surfaces, lines body cavities, "
                  "and forms glands. It is the most widely distributed tissue in the body. "
                  "The term 'epithelium' was coined by Frederik Ruysch.")

    # General Characteristics
    add_heading_styled(doc, "1.1 General Characteristics", 2)
    add_bullet_points(doc, [
        "Cellularity: Composed almost entirely of cells with minimal extracellular matrix",
        "Polarity: Has apical (free) surface and basal surface - structural and functional differences",
        "Attachment: Connected to underlying connective tissue by basement membrane",
        "Avascularity: No blood vessels; nutrients diffuse from underlying connective tissue",
        "Innervation: Has nerve supply (sensory nerve endings present)",
        "Regeneration: High mitotic rate; capable of rapid renewal and repair",
        "Specialized contacts: Cells held together by tight junctions, desmosomes, gap junctions",
        "Surface specializations: Microvilli, cilia, stereocilia on apical surface"
    ], bold_prefix=True)

    doc.add_paragraph()
    add_highlight_box(doc, "HIGH YIELD: Basement Membrane",
                     ["Composed of: Basal lamina + Reticular lamina",
                      "Basal lamina = Lamina lucida (clear) + Lamina densa (dark)",
                      "Made of: Type IV collagen, Laminin, Fibronectin, Proteoglycans",
                      "Functions: Structural support, Molecular filter, Cell guidance during repair",
                      "Clinical: Thickened in Diabetes Mellitus (esp. kidney glomerulus)"],
                     color="E8F8F5", border_color="1ABC9C")

    # Classification
    doc.add_paragraph()
    add_heading_styled(doc, "1.2 Classification of Epithelial Tissue", 2)
    add_body_text(doc, "Classification is based on: (A) Number of cell layers, and (B) Shape of cells at the free surface.")

    add_heading_styled(doc, "A. Based on Number of Layers:", 3)
    add_bullet_points(doc, [
        "Simple epithelium: Single layer of cells (all cells rest on basement membrane)",
        "Stratified epithelium: Two or more layers (only basal layer touches basement membrane)",
        "Pseudostratified epithelium: Appears stratified but all cells rest on basement membrane (not all reach surface)"
    ], bold_prefix=True)

    add_heading_styled(doc, "B. Based on Cell Shape:", 3)
    add_bullet_points(doc, [
        "Squamous: Flat, scale-like cells (wider than tall)",
        "Cuboidal: Cube-shaped cells (equal dimensions)",
        "Columnar: Column-shaped cells (taller than wide)",
        "Transitional: Changeable shape (dome-shaped when relaxed, flat when stretched)"
    ], bold_prefix=True)

    doc.add_paragraph()

    # Detailed types table
    add_heading_styled(doc, "1.3 Types of Simple Epithelium", 2)

    create_styled_table(doc,
        ["Type", "Structure", "Location", "Function"],
        [
            ["Simple Squamous", "Single layer of flat cells\nNucleus: central, disc-shaped",
             "Alveoli of lungs\nLining of blood vessels (endothelium)\nLining of body cavities (mesothelium)\nBowman's capsule\nLoop of Henle (thin segment)",
             "Diffusion\nFiltration\nOsmosis\nReduces friction"],
            ["Simple Cuboidal", "Single layer of cube-shaped cells\nNucleus: central, round",
             "Kidney tubules (PCT, DCT)\nThyroid follicles\nOvary surface (germinal epithelium)\nLens of eye\nSmall ducts of glands",
             "Secretion\nAbsorption\nProtection"],
            ["Simple Columnar\n(Non-ciliated)", "Single layer of tall cells\nNucleus: basal, oval\nGoblet cells present",
             "Stomach to anus (GIT)\nGall bladder\nLarge ducts of glands",
             "Secretion (mucus)\nAbsorption\nProtection"],
            ["Simple Columnar\n(Ciliated)", "Columnar cells with cilia\nGoblet cells present",
             "Uterine tubes (Fallopian)\nSmall bronchi\nUterus\nCentral canal of spinal cord",
             "Movement of ovum\nMovement of mucus"],
            ["Pseudostratified\nColumnar (Ciliated)", "All cells on BM, not all reach surface\nGoblet cells, cilia present",
             "Trachea & bronchi\nNasal cavity\nAuditory tube (Eustachian)\nMale urethra (parts)",
             "Secretion of mucus\nMovement of mucus\n(Mucociliary escalator)"],
        ],
        header_color="1A5276")

    doc.add_paragraph()
    add_highlight_box(doc, "IMPORTANT: Special Epithelia Names",
                     ["Endothelium: Simple squamous epithelium lining blood vessels, lymphatics, heart",
                      "Mesothelium: Simple squamous epithelium lining body cavities (pleural, peritoneal, pericardial)",
                      "Neuroepithelium: Specialized sensory epithelium (retina, organ of Corti, olfactory mucosa)",
                      "Myoepithelium: Contractile epithelial cells around glands (salivary, mammary, sweat glands)",
                      "Germinal epithelium: Cuboidal epithelium covering ovary surface"],
                     color="FDEDEC", border_color="E74C3C")

    doc.add_page_break()


    # Stratified Epithelium
    add_heading_styled(doc, "1.4 Types of Stratified Epithelium", 2)

    create_styled_table(doc,
        ["Type", "Structure", "Location", "Function"],
        [
            ["Stratified Squamous\n(Keratinized)", "Multiple layers\nSurface cells flat & dead\nFilled with keratin\nNo nuclei in surface cells",
             "Skin (epidermis)\nLips (outer)\nHard palate",
             "Protection against\nabrasion, water loss,\nchemicals, pathogens"],
            ["Stratified Squamous\n(Non-keratinized)", "Multiple layers\nSurface cells flat but alive\nNuclei present\nMoist surface",
             "Oral cavity\nEsophagus\nVagina\nAnal canal\nCornea",
             "Protection against\nabrasion\nMoist areas"],
            ["Stratified Cuboidal", "2-3 layers of cuboidal cells\nRelatively rare",
             "Ducts of sweat glands\nDucts of mammary glands\nPart of male urethra",
             "Protection\nLimited secretion"],
            ["Stratified Columnar", "Basal layers: cuboidal/columnar\nSurface: columnar\nRare tissue",
             "Parts of pharynx\nMale urethra\nLarge ducts (parotid)\nConjunctiva of eye",
             "Protection\nSecretion"],
            ["Transitional\n(Urothelium)", "Relaxed: dome-shaped cells (6-7 layers)\nStretched: flat cells (2-3 layers)\nUmbrella cells on surface",
             "Urinary bladder\nUreter\nRenal pelvis\nUpper urethra",
             "Stretching\n(distensibility)\nProtection from\nurine toxicity"],
        ],
        header_color="6C3483")

    doc.add_paragraph()
    add_highlight_box(doc, "CLINICAL CORRELATION: Metaplasia",
                     ["Metaplasia = Reversible change from one epithelial type to another",
                      "Example: Smokers - Pseudostratified ciliated columnar of trachea changes to stratified squamous",
                      "Barrett's esophagus: Stratified squamous changes to simple columnar (due to acid reflux)",
                      "Vitamin A deficiency: Columnar epithelium changes to keratinized squamous",
                      "Metaplasia is reversible if stimulus is removed; Dysplasia may follow if stimulus persists"],
                     color="F9EBEA", border_color="CB4335")

    # Cell Junctions
    doc.add_paragraph()
    add_heading_styled(doc, "1.5 Cell Junctions (Intercellular Junctions)", 2)

    add_body_text(doc, "Cell junctions are specialized structures that connect epithelial cells "
                  "to each other and to the basement membrane. They form the Junctional Complex.")

    create_styled_table(doc,
        ["Junction Type", "Structure", "Function", "Location"],
        [
            ["Tight Junctions\n(Zonula Occludens)", "Claudins & Occludins\nCompletely encircle cell\nFusion of outer membrane leaflets",
             "Seal between cells\nPrevent paracellular passage\nMaintain polarity",
             "Intestinal epithelium\nBBB (brain)\nRenal tubules"],
            ["Adherens Junctions\n(Zonula Adherens)", "Cadherins (E-cadherin)\nLinked to actin filaments\nCalcium-dependent",
             "Cell-cell adhesion\nSignal transduction\nMorphogenesis",
             "All epithelia\nCardiac muscle\n(fascia adherens)"],
            ["Desmosomes\n(Macula Adherens)", "Desmogleins & Desmocollins\nLinked to intermediate filaments (keratin)\nButton-like spots",
             "Strong cell-cell adhesion\nResist mechanical stress\n'Spot welds'",
             "Skin (epidermis)\nCardiac muscle\nCervix"],
            ["Gap Junctions\n(Nexus/Communicating)", "Connexins form connexons\n6 connexins = 1 connexon\nChannel diameter: 1.5-2nm",
             "Direct cell communication\nIon & small molecule passage\nElectrical coupling",
             "Cardiac muscle\nSmooth muscle\nNerve synapses"],
            ["Hemidesmosomes", "Integrins (not cadherins)\nLinked to intermediate filaments\nAttach to BM",
             "Anchor cells to BM\nResist shearing forces",
             "Epidermal basal cells\nCorneal epithelium"],
        ],
        header_color="1E8449")

    doc.add_paragraph()
    add_highlight_box(doc, "CLINICAL: Diseases of Cell Junctions",
                     ["Pemphigus vulgaris: Autoantibodies against desmogleins (desmosomes) - skin blisters",
                      "Bullous pemphigoid: Autoantibodies against hemidesmosomes - subepidermal blisters",
                      "Epidermolysis bullosa: Genetic defect in hemidesmosomes/anchoring fibrils",
                      "Hereditary hemorrhagic telangiectasia: Defect in gap junctions"],
                     color="FEF9E7", border_color="F39C12")

    doc.add_page_break()


    # Surface Modifications
    add_heading_styled(doc, "1.6 Surface Modifications of Epithelial Cells", 2)

    create_styled_table(doc,
        ["Modification", "Structure", "Location", "Function"],
        [
            ["Microvilli", "Finger-like projections\nActin core (20-30 filaments)\n1-2 um long\nForm brush/striated border",
             "Small intestine (striated border)\nKidney PCT (brush border)\nGall bladder",
             "Increase surface area\nfor absorption\n(20-30x increase)"],
            ["Cilia", "Longer than microvilli (5-10 um)\n9+2 microtubule arrangement\nDynein arms (motor protein)\nPowered by ATP",
             "Trachea, bronchi\nFallopian tubes\nVentricles of brain\nMiddle ear",
             "Movement of substances\nMucociliary clearance\nOvum transport"],
            ["Stereocilia", "Long, non-motile microvilli\n(NOT true cilia - no 9+2)\nBranched, lack dynein",
             "Epididymis\nVas deferens\nHair cells of inner ear",
             "Absorption\nMechanotransduction\n(hearing, balance)"],
            ["Flagella", "Single, long (150 um)\nSame 9+2 structure as cilia\nWhip-like movement",
             "Spermatozoa (tail)",
             "Cell locomotion"],
        ],
        header_color="2C3E50")

    doc.add_paragraph()
    add_highlight_box(doc, "HIGH YIELD: Ciliary Defects",
                     ["Kartagener syndrome (Primary Ciliary Dyskinesia):",
                      "  - Defect in dynein arms of cilia",
                      "  - Triad: Situs inversus + Bronchiectasis + Sinusitis",
                      "  - Also causes male infertility (immotile sperm)",
                      "  - Autosomal recessive inheritance"],
                     color="E8F8F5", border_color="16A085")

    # Glandular Epithelium
    doc.add_paragraph()
    add_heading_styled(doc, "1.7 Glandular Epithelium", 2)

    add_body_text(doc, "Glands are formed by infolding of epithelial cells. A gland is one or more "
                  "cells that secrete a particular product (secretion).")

    add_heading_styled(doc, "Classification of Glands:", 3)

    create_styled_table(doc,
        ["Basis", "Type", "Description", "Examples"],
        [
            ["Destination of\nSecretion", "Exocrine", "Secrete into ducts\nor onto surfaces", "Salivary, Sweat,\nSebaceous, Pancreas (exo)"],
            ["", "Endocrine", "Secrete into blood\n(ductless glands)", "Thyroid, Adrenal,\nPituitary, Pancreas (endo)"],
            ["", "Mixed/Heterocrine", "Both exocrine &\nendocrine function", "Pancreas, Liver,\nGonads"],
            ["Cell Number", "Unicellular", "Single-celled gland", "Goblet cells"],
            ["", "Multicellular", "Many cells form gland", "Salivary glands, Liver"],
            ["Mode of Secretion", "Merocrine (Eccrine)", "Exocytosis\nCell intact", "Salivary glands,\nPancreas, Sweat glands"],
            ["", "Apocrine", "Apical portion pinched off\nwith secretion", "Mammary gland,\nApocrine sweat glands"],
            ["", "Holocrine", "Entire cell ruptures\nand dies to release", "Sebaceous glands"],
            ["Type of Secretion", "Serous", "Watery, enzyme-rich", "Parotid gland"],
            ["", "Mucous", "Thick, viscous mucin", "Sublingual gland"],
            ["", "Mixed (Seromucous)", "Both serous and mucous", "Submandibular gland"],
        ],
        header_color="7D3C98")

    doc.add_paragraph()
    add_highlight_box(doc, "MNEMONIC: Mode of Secretion - 'MAH'",
                     ["Merocrine - Most glands (exocytosis, cell intact)",
                      "Apocrine - Apex lost (mammary, some sweat glands)",
                      "Holocrine - Hole cell dies (sebaceous glands)",
                      "Remember: Merocrine > Apocrine > Holocrine (increasing cell destruction)"],
                     color="EBF5FB", border_color="2E86C1")

    doc.add_page_break()


def add_connective_tissue(doc):
    """Add comprehensive connective tissue section."""
    add_heading_styled(doc, "SECTION 2: CONNECTIVE TISSUE", 1)
    add_section_divider(doc)

    add_body_text(doc, "Connective tissue is the most abundant and widely distributed tissue in the body. "
                  "It connects, supports, and binds other tissues. Unlike epithelium, it is characterized by "
                  "large amounts of extracellular matrix with relatively few cells scattered within it.")

    # General Characteristics
    add_heading_styled(doc, "2.1 General Characteristics", 2)
    add_bullet_points(doc, [
        "Origin: Derived from mesoderm (mesenchyme - embryonic connective tissue)",
        "Vascularity: Most are highly vascular (exception: cartilage, tendons, ligaments)",
        "Matrix: Abundant extracellular matrix (ECM) - key distinguishing feature",
        "Components: Cells + Fibers + Ground substance = Extracellular Matrix",
        "Functions: Support, Protection, Transport, Storage, Defense, Repair"
    ], bold_prefix=True)

    # Components
    doc.add_paragraph()
    add_heading_styled(doc, "2.2 Components of Connective Tissue", 2)

    add_heading_styled(doc, "A. Cells of Connective Tissue", 3)

    create_styled_table(doc,
        ["Cell Type", "Origin", "Function", "Special Features"],
        [
            ["Fibroblasts", "Mesenchyme", "Synthesize fibers & ground substance\nMost common CT cell\nWound repair",
             "Active: large, pale nucleus\nInactive (fibrocyte): small, dark\nMost numerous in CT proper"],
            ["Macrophages\n(Histiocytes)", "Monocytes (blood)", "Phagocytosis\nAntigen presentation\nCytokine secretion",
             "Fixed (resident) or Wandering\nPart of MPS system\nDust cells (lung)\nKupffer cells (liver)\nMicroglia (brain)\nOsteoclasts (bone)"],
            ["Mast Cells", "Bone marrow\n(CD34+ precursors)", "Release histamine & heparin\nAllergic reactions\nInflammation",
             "Metachromatic granules\nToluidine blue staining\nIgE receptors (FcERI)\nTryptase = specific marker"],
            ["Plasma Cells", "B-lymphocytes", "Antibody (Ig) production\nHumoral immunity",
             "Clock-face/cartwheel nucleus\nAbundant RER\nPerinuclear halo (Golgi)\nFound in GIT, respiratory tract"],
            ["Adipocytes", "Mesenchyme\n(lipoblasts)", "Energy storage (fat)\nInsulation\nCushioning\nEndocrine functions",
             "White fat: single large droplet\n(unilocular, signet ring)\nBrown fat: multiple droplets\n(multilocular, in infants)"],
            ["Mesenchymal\nStem Cells", "Mesenchyme", "Differentiate into other CT cells\nRepair & regeneration",
             "Undifferentiated\nRetain mitotic ability\nFound along blood vessels"],
            ["Leukocytes\n(WBCs)", "Bone marrow", "Defense & immunity\nInflammation",
             "Neutrophils (acute inflammation)\nLymphocytes (chronic)\nEosinophils (parasites, allergy)"],
        ],
        header_color="1A5276")

    doc.add_paragraph()
    add_highlight_box(doc, "MONONUCLEAR PHAGOCYTE SYSTEM (MPS)",
                     ["Previously called Reticuloendothelial System (RES)",
                      "All derived from monocytes:",
                      "  Blood: Monocytes",
                      "  Liver: Kupffer cells",
                      "  Lung: Alveolar macrophages (Dust cells)",
                      "  Brain: Microglia",
                      "  Bone: Osteoclasts",
                      "  Skin: Langerhans cells (also dendritic cells)",
                      "  Spleen: Splenic macrophages",
                      "  Kidney: Mesangial cells",
                      "  Lymph node: Sinus histiocytes"],
                     color="F4ECF7", border_color="8E44AD")

    doc.add_paragraph()

    add_heading_styled(doc, "B. Fibers of Connective Tissue", 3)

    create_styled_table(doc,
        ["Fiber Type", "Composition", "Properties", "Location", "Staining"],
        [
            ["Collagen\n(White fibers)", "Type I (90%): skin, bone, tendon\nType II: cartilage\nType III: reticular\nType IV: basement membrane\n28+ types known",
             "Great tensile strength\nFlexible but inextensible\nMost abundant protein\nin body (25-30%)",
             "Skin, bone, tendon\nLigaments, dentin\nCartilage, BM\nAll CT types",
             "Eosin (pink) - H&E\nVan Gieson: Red\nMasson trichrome: Blue/Green"],
            ["Elastic\n(Yellow fibers)", "Elastin core +\nFibrillin microfibrils\n(Fibrillin-1 glycoprotein)",
             "Stretch & recoil (150%)\nBranching fibers\nLess tensile strength\nthan collagen",
             "Lungs (alveolar walls)\nAorta & large arteries\nSkin (dermis)\nEar pinna\nEpiglottis\nLigamentum flavum\nVocal cords",
             "Orcein stain: Brown\nVerhoff stain: Black\nAldehyde fuchsin: Purple"],
            ["Reticular\n(Argyrophilic)", "Type III collagen\nThin, branching network\nCoated with glycoproteins",
             "Form framework/stroma\nSupport other cells\nDelicate scaffolding",
             "Lymph nodes, Spleen\nBone marrow\nLiver (around sinusoids)\nBasement membrane",
             "Silver stains: Black\n(argyrophilic)\nPAS: Positive"],
        ],
        header_color="154360")

    doc.add_paragraph()
    add_highlight_box(doc, "CLINICAL: Collagen & Elastic Fiber Diseases",
                     ["Scurvy: Vitamin C deficiency - defective collagen synthesis (hydroxylation of proline/lysine)",
                      "Ehlers-Danlos syndrome: Defective collagen synthesis - hyperextensible skin, hypermobile joints",
                      "Osteogenesis imperfecta: Defective Type I collagen - brittle bone disease",
                      "Marfan syndrome: Defect in Fibrillin-1 (elastic fibers) - tall stature, lens dislocation, aortic aneurysm",
                      "Alport syndrome: Defect in Type IV collagen (BM) - glomerulonephritis, hearing loss"],
                     color="FDEDEC", border_color="E74C3C")

    doc.add_page_break()


    # Ground Substance
    add_heading_styled(doc, "C. Ground Substance", 3)

    add_body_text(doc, "Ground substance is the gel-like material that fills the space between cells "
                  "and fibers. It is transparent, colorless, and homogeneous.")

    add_bullet_points(doc, [
        "Glycosaminoglycans (GAGs): Hyaluronic acid, Chondroitin sulfate, Keratan sulfate, Heparan sulfate, Dermatan sulfate",
        "Proteoglycans: Core protein + GAG chains (e.g., Aggrecan in cartilage)",
        "Glycoproteins: Fibronectin (binds cells to matrix), Laminin (in BM), Chondronectin, Osteonectin",
        "Water: Major component; provides medium for diffusion of nutrients and wastes",
        "Properties: Acts as molecular sieve, barrier to bacteria, lubricant"
    ], bold_prefix=True)

    # Classification of CT
    doc.add_paragraph()
    add_heading_styled(doc, "2.3 Classification of Connective Tissue", 2)

    add_heading_styled(doc, "Overview of CT Classification:", 3)
    add_numbered_list(doc, [
        "Embryonic CT: Mesenchyme, Mucous CT (Wharton's jelly)",
        "CT Proper: (a) Loose/Areolar, (b) Dense - Regular & Irregular, (c) Elastic, (d) Reticular, (e) Adipose",
        "Specialized/Supporting CT: (a) Cartilage, (b) Bone, (c) Blood"
    ])

    doc.add_paragraph()
    add_heading_styled(doc, "2.4 Connective Tissue Proper", 2)

    create_styled_table(doc,
        ["Type", "Description", "Location", "Function"],
        [
            ["Loose (Areolar) CT", "Most common CT type\nAll 3 fiber types present\nMany cell types\nSemifluid ground substance",
             "Beneath epithelium (lamina propria)\nAround blood vessels & nerves\nBetween muscles\nSubcutaneous tissue (superficial fascia)",
             "Wrapping & cushioning\nHolds tissue fluid\nDefense (immune cells)\nNourishes epithelium"],
            ["Dense Regular CT", "Parallel collagen bundles\nFew cells (fibroblasts in rows)\nMinimal ground substance\nGreat tensile strength",
             "Tendons (muscle to bone)\nLigaments (bone to bone)\nAponeuroses",
             "Resist pulling force\nin one direction\nConnect structures"],
            ["Dense Irregular CT", "Collagen in random directions\nFibroblasts scattered\nThick bundles of collagen",
             "Dermis of skin (reticular layer)\nJoint capsules\nPeriosteum, Perichondrium\nOrgan capsules\nHeart valves\nSubmucosa of GIT",
             "Resist tension in\nmultiple directions\nProvide strength"],
            ["Elastic CT", "Predominantly elastic fibers\nBranching elastic sheets\nYellow color",
             "Walls of large arteries (aorta)\nLigamentum flavum (spine)\nTrachea & bronchi\nTrue vocal cords\nLigamentum nuchae",
             "Allow stretch & recoil\nMaintain blood pressure\nSpring-back mechanism"],
            ["Reticular CT", "Reticular fibers (Type III collagen)\nReticular cells\nDelicate meshwork",
             "Stroma of: Spleen, Lymph nodes,\nBone marrow, Liver\nBasement membrane",
             "Support framework\nFilter function\nHouse immune cells"],
            ["Adipose Tissue\n(White)", "Unilocular fat cells\nSignet ring appearance\nRichly vascularized",
             "Subcutaneous layer\nAround kidneys (perirenal)\nBehind eyes (orbital)\nMesentery, Omentum\nYellow bone marrow",
             "Energy storage\nInsulation\nCushioning/protection\nEndocrine (leptin)"],
            ["Adipose Tissue\n(Brown)", "Multilocular fat cells\nAbundant mitochondria\n(give brown color)\nUncoupling Protein-1 (UCP-1/thermogenin)",
             "Newborns: interscapular region\nAround great vessels\nAdrenal, kidney hilum\nAdults: minimal",
             "Heat production\n(non-shivering\nthermogenesis)\nNeonatal thermoregulation"],
        ],
        header_color="1E8449")

    doc.add_paragraph()
    add_highlight_box(doc, "HIGH YIELD: Adipose Tissue Facts",
                     ["White adipose (WAT): Energy storage, insulation, leptin secretion",
                      "Brown adipose (BAT): Thermogenesis via UCP-1 (uncouples oxidative phosphorylation)",
                      "BAT has more mitochondria, more blood supply, more sympathetic innervation",
                      "Lipoma: Benign tumor of adipose tissue (most common soft tissue tumor)",
                      "Liposarcoma: Malignant tumor of adipose tissue"],
                     color="E8F8F5", border_color="1ABC9C")

    doc.add_page_break()


    # Cartilage
    add_heading_styled(doc, "2.5 Cartilage (Supporting CT)", 2)

    add_body_text(doc, "Cartilage is a specialized connective tissue with a firm, flexible matrix. "
                  "It is avascular, aneural, and alymphatic. Nutrients reach chondrocytes by diffusion "
                  "through the matrix from surrounding perichondrium.")

    add_heading_styled(doc, "General Features of Cartilage:", 3)
    add_bullet_points(doc, [
        "Cells: Chondroblasts (active), Chondrocytes (mature, in lacunae)",
        "Matrix: Firm but flexible; rich in Type II collagen + proteoglycans (aggrecan)",
        "Perichondrium: Dense CT covering (absent in articular, fibrocartilage)",
        "Growth: Interstitial (from within) + Appositional (from perichondrium)",
        "Avascular: Nutrients by diffusion only (limits thickness to ~3mm)",
        "Repair: Poor regeneration capacity (except in children)"
    ], bold_prefix=True)

    doc.add_paragraph()

    create_styled_table(doc,
        ["Type", "Matrix", "Location", "Features"],
        [
            ["Hyaline Cartilage\n(Most common)", "Type II collagen (fine)\nProteoglycans (aggrecan)\nGlassy, bluish-white\nChondroitin sulfate",
             "Articular surfaces\nTracheal & bronchial rings\nNasal septum\nCostal cartilages\nFetal skeleton\nLarynx (thyroid, cricoid)\nEpiphyseal plates",
             "Most abundant type\nForms model for bone (endochondral)\nFreezing point: -20C\nPearly blue appearance\nIsogenous groups (cell nests)"],
            ["Elastic Cartilage", "Type II collagen +\nAbundant elastic fibers\nYellow color\n(Orcein/Verhoeff stains elastic fibers)",
             "Ear pinna (auricle)\nEpiglottis\nAuditory tube (Eustachian)\nCorniculate & Cuneiform cartilages\nExternal ear canal",
             "More flexible than hyaline\nNEVER calcifies\nAlways has perichondrium\nYellowish appearance\nMaintains shape with flexibility"],
            ["Fibrocartilage", "Type I collagen (dense, thick)\n+ Type II collagen\nLess ground substance\nRows of chondrocytes",
             "Intervertebral discs\nPublic symphysis\nMenisci of knee\nTMJ disc\nLabrum (glenoid, acetabular)\nInsertion of tendons into bone",
             "Strongest cartilage\nNo perichondrium\nTransitional between\ndense CT and hyaline\nCombines strength with\nsome compressibility"],
        ],
        header_color="7D3C98")

    doc.add_paragraph()
    add_highlight_box(doc, "CLINICAL: Cartilage Disorders",
                     ["Achondroplasia: Defect in FGFR3 gene - dwarfism (most common cause)",
                      "Herniated disc: Nucleus pulposus protrudes through annulus fibrosus (fibrocartilage)",
                      "Osteoarthritis: Degeneration of articular (hyaline) cartilage",
                      "Relapsing polychondritis: Autoimmune destruction of cartilage (ear, nose, trachea)",
                      "Chondrosarcoma: Malignant tumor of cartilage"],
                     color="FEF9E7", border_color="F39C12")

    doc.add_page_break()


    # Bone
    add_heading_styled(doc, "2.6 Bone (Osseous Tissue)", 2)

    add_body_text(doc, "Bone is the hardest connective tissue. It provides structural support, "
                  "protection, mineral storage, and houses bone marrow for hematopoiesis.")

    add_heading_styled(doc, "General Features:", 3)
    add_bullet_points(doc, [
        "Cells: Osteoblasts (bone forming), Osteocytes (mature, in lacunae), Osteoclasts (bone resorbing - multinucleated, from monocytes)",
        "Matrix: Organic (35%): Type I collagen + Ground substance | Inorganic (65%): Hydroxyapatite crystals [Ca10(PO4)6(OH)2]",
        "Periosteum: Outer fibrous layer + Inner osteogenic (cambium) layer",
        "Endosteum: Lines marrow cavity; contains osteoprogenitor cells",
        "Blood supply: Highly vascular (unlike cartilage)",
        "Haversian system (Osteon): Structural unit of compact bone"
    ], bold_prefix=True)

    doc.add_paragraph()

    create_styled_table(doc,
        ["Feature", "Compact (Dense/Cortical) Bone", "Spongy (Cancellous/Trabecular) Bone"],
        [
            ["Structure", "Dense, solid; organized into osteons (Haversian systems)", "Lattice of thin plates (trabeculae) with marrow spaces"],
            ["Location", "Outer layer of all bones; shaft (diaphysis) of long bones", "Interior of bones; epiphyses of long bones; flat bones"],
            ["Osteons", "Present (concentric lamellae around Haversian canal)", "Absent; irregular lamellae in trabeculae"],
            ["Blood Supply", "Through Haversian & Volkmann canals", "From marrow cavity directly"],
            ["Marrow", "Yellow marrow (fat) in medullary cavity", "Red marrow (hematopoietic) in spaces"],
            ["Weight", "~80% of skeletal mass", "~20% of skeletal mass"],
            ["Function", "Strength, support, protection", "Lightweight support; hematopoiesis; mineral exchange"],
        ],
        header_color="1A5276")

    doc.add_paragraph()
    add_heading_styled(doc, "Bone Formation (Ossification):", 3)
    add_numbered_list(doc, [
        "Intramembranous ossification: Direct bone formation from mesenchyme (flat bones of skull, clavicle, mandible)",
        "Endochondral ossification: Bone replaces hyaline cartilage model (most bones - long bones, vertebrae, pelvis)"
    ])

    doc.add_paragraph()
    add_highlight_box(doc, "CLINICAL: Bone Disorders",
                     ["Osteoporosis: Decreased bone mass & density; fracture risk (postmenopausal women)",
                      "Osteomalacia/Rickets: Defective mineralization (Vitamin D deficiency); soft bones",
                      "Paget disease: Excessive bone remodeling; mosaic pattern of bone",
                      "Osteogenesis imperfecta: Type I collagen defect; brittle bones; blue sclera",
                      "Osteopetrosis: Defective osteoclasts; dense but brittle bones (marble bone disease)"],
                     color="F9EBEA", border_color="CB4335")

    # Blood as CT
    doc.add_paragraph()
    add_heading_styled(doc, "2.7 Blood (Fluid Connective Tissue)", 2)

    add_body_text(doc, "Blood is a specialized connective tissue with a liquid extracellular matrix "
                  "called plasma. It transports nutrients, gases, hormones, and waste products.")

    add_bullet_points(doc, [
        "Matrix: Plasma (55%) - water, proteins (albumin, globulins, fibrinogen), electrolytes",
        "Formed elements (45%): RBCs (erythrocytes), WBCs (leukocytes), Platelets (thrombocytes)",
        "RBCs: Biconcave disc, no nucleus, contain hemoglobin, 120-day lifespan, 5 million/uL",
        "WBCs: 5000-10000/uL; Granulocytes (Neutrophils, Eosinophils, Basophils) + Agranulocytes (Lymphocytes, Monocytes)",
        "Platelets: Cell fragments from megakaryocytes, 2.5 lakh/uL, hemostasis",
        "Hematopoiesis: Formation of blood cells in bone marrow from pluripotent stem cells"
    ], bold_prefix=True)

    doc.add_paragraph()
    add_highlight_box(doc, "MNEMONIC: WBC Types (Decreasing Order) - 'Never Let Monkeys Eat Bananas'",
                     ["Neutrophils (60-70%) > Lymphocytes (20-25%) > Monocytes (3-8%) > Eosinophils (2-4%) > Basophils (0.5-1%)",
                      "Alternative: 'Neutrophils Love Making Everything Better'"],
                     color="EBF5FB", border_color="2E86C1")

    doc.add_page_break()


def add_muscular_tissue(doc):
    """Add comprehensive muscular tissue section."""
    add_heading_styled(doc, "SECTION 3: MUSCULAR TISSUE", 1)
    add_section_divider(doc)

    add_body_text(doc, "Muscular tissue is responsible for body movements. It is characterized by the "
                  "properties of contractility, excitability, extensibility, and elasticity. "
                  "Muscle cells (fibers) contain specialized contractile proteins - actin and myosin.")

    # General Properties
    add_heading_styled(doc, "3.1 General Properties of Muscle Tissue", 2)
    add_bullet_points(doc, [
        "Excitability (Irritability): Ability to respond to stimuli (electrical, chemical, mechanical)",
        "Contractility: Ability to shorten and generate force when stimulated",
        "Extensibility: Ability to be stretched beyond resting length without damage",
        "Elasticity: Ability to return to original shape after being stretched",
        "Conductivity: Ability to propagate electrical impulses along the cell membrane"
    ], bold_prefix=True)

    # Special Terminology
    doc.add_paragraph()
    add_heading_styled(doc, "3.2 Muscle Terminology (Unique to Muscle Tissue)", 2)

    create_styled_table(doc,
        ["General Term", "Muscle Term", "Meaning"],
        [
            ["Cell/Fiber", "Muscle fiber (Myofiber)", "Individual muscle cell"],
            ["Cell membrane", "Sarcolemma", "Plasma membrane of muscle cell"],
            ["Cytoplasm", "Sarcoplasm", "Cytoplasm of muscle cell"],
            ["Endoplasmic reticulum", "Sarcoplasmic reticulum (SR)", "Stores calcium ions"],
            ["Mitochondria", "Sarcosomes", "Energy production (ATP)"],
            ["Bundle of fibers", "Fasciculus (Fascicle)", "Group of muscle fibers"],
        ],
        header_color="2C3E50")

    # Classification - Main Comparison Table
    doc.add_paragraph()
    add_heading_styled(doc, "3.3 Types of Muscle Tissue - Comprehensive Comparison", 2)

    create_styled_table(doc,
        ["Feature", "Skeletal Muscle", "Cardiac Muscle", "Smooth Muscle"],
        [
            ["Alternative names", "Voluntary, Striated,\nSomatic", "Involuntary, Striated,\nCardiac", "Involuntary, Non-striated,\nVisceral, Plain"],
            ["Location", "Attached to skeleton\nTongue, Pharynx\nUpper esophagus\nDiaphragm, Eye muscles", "Heart (myocardium)\nPulmonary veins\nSuperior vena cava", "Walls of hollow organs\n(GIT, urinary, reproductive)\nBlood vessels, Airways\nIris, Arrector pili"],
            ["Shape", "Long, cylindrical\n(1-40 cm length)\nUnbranched", "Short, branching\nY-shaped\nIrregular cylinders", "Spindle-shaped (fusiform)\nSmall (20-500 um)\nPointed ends"],
            ["Nuclei", "Multiple (multinucleated)\nPeripheral (beneath sarcolemma)\n(formed by myoblast fusion)", "1-2 nuclei (usually 1)\nCentral location\nOval shape", "Single nucleus\nCentral location\nCigar/elongated shape"],
            ["Striations", "Present (prominent)\nA bands, I bands, Z lines\nH zone, M line", "Present (less prominent)\nSame banding pattern\nas skeletal", "Absent\nNo organized sarcomere\nDense bodies instead of Z lines"],
            ["Sarcomere", "Present (2.2 um)\nFunctional unit of contraction", "Present\nSimilar to skeletal", "Absent\nDense bodies serve as\nattachment points"],
            ["T-tubules", "Present at A-I junction\nTriads (1 T-tubule + 2 SR)", "Present at Z-line\nDyads (1 T-tubule + 1 SR)\nLarger than skeletal", "Absent\nCaveolae instead\n(membrane invaginations)"],
            ["Intercalated discs", "Absent", "PRESENT (unique feature)\nContain: desmosomes +\ngap junctions +\nfascia adherens", "Absent\nGap junctions present\n(but no intercalated discs)"],
            ["Motor unit", "Present\n(1 neuron: few to 1000+ fibers)", "Absent\n(functional syncytium)", "Absent\n(functional syncytium)"],
            ["Neuromuscular\njunction", "Present (motor end plate)\nACh is neurotransmitter\nNicotinic receptors", "No NMJ\nAuto-rhythmic\n(SA node pacemaker)", "Varicosities (no NMJ)\nMultiple neurotransmitters\n(ACh, NE)"],
            ["Control", "Voluntary\n(somatic nervous system)", "Involuntary\n(ANS modulates rate)\nIntrinsic rhythm", "Involuntary\n(ANS, hormones,\nlocal factors)"],
            ["Contraction speed", "Fast (but varies)\nFast-twitch (Type II)\nSlow-twitch (Type I)", "Moderate\nRhythmic\n72 bpm (resting)", "Slowest\nSustained (tonic)\nRhythmic in GIT"],
            ["Fatigue", "Fatigues relatively quickly\n(except Type I fibers)", "Fatigue resistant\n(abundant mitochondria\n& myoglobin)", "Very fatigue resistant\n(low energy requirement)"],
            ["Regeneration", "Limited\n(satellite cells)\nHypertrophy > hyperplasia", "Minimal/None\nDamage replaced by scar\n(fibrosis)", "Good regenerative capacity\nCan undergo hyperplasia\n(e.g., pregnant uterus)"],
            ["Ca2+ source", "Sarcoplasmic reticulum\n(intracellular only)", "SR + Extracellular\n(Ca-induced Ca release)\nDHP + Ryanodine receptors", "Extracellular (mainly) +\nSR (some)\nCalmodulin-dependent"],
            ["Regulatory protein", "Troponin-Tropomyosin\n(on thin filament)", "Troponin-Tropomyosin\n(same as skeletal)", "Calmodulin\n(activates MLCK)\nNo troponin"],
        ],
        header_color="1B4F72")

    doc.add_page_break()


    # Skeletal Muscle Detail
    add_heading_styled(doc, "3.4 Skeletal Muscle - Detailed Structure", 2)

    add_heading_styled(doc, "Organization (Outer to Inner):", 3)
    add_numbered_list(doc, [
        "Epimysium: Dense irregular CT covering entire muscle (deep fascia)",
        "Perimysium: Surrounds each fascicle (bundle of fibers)",
        "Endomysium: Delicate reticular fiber around each individual muscle fiber",
        "Sarcolemma: Cell membrane of individual muscle fiber",
        "Sarcoplasm: Contains myofibrils, mitochondria, glycogen, myoglobin"
    ])

    doc.add_paragraph()
    add_heading_styled(doc, "Sarcomere Structure (Functional Unit):", 3)
    add_body_text(doc, "The sarcomere extends from one Z-line to the next Z-line (~2.2 um at rest):")

    create_styled_table(doc,
        ["Band/Line", "Composition", "Behavior During Contraction"],
        [
            ["A Band (Anisotropic)", "Dark band; contains thick (myosin) filaments\n+ overlap of thin filaments at edges", "Remains CONSTANT in length"],
            ["I Band (Isotropic)", "Light band; thin (actin) filaments only\nBisected by Z-line", "SHORTENS during contraction"],
            ["H Zone (Helle)", "Center of A band; thick filaments only\n(no overlap with thin)", "SHORTENS/disappears"],
            ["Z Line (Disc/Zwischenscheibe)", "Anchoring point for thin filaments\nAlpha-actinin protein", "Z-lines come CLOSER together"],
            ["M Line (Mittelscheibe)", "Center of H zone; holds thick filaments\nMyomesin, C-protein", "Remains in center"],
        ],
        header_color="6C3483")

    doc.add_paragraph()
    add_highlight_box(doc, "MNEMONIC: Sliding Filament Theory",
                     ["During contraction: 'HAI' decreases (H zone, A-I junction, I band all shorten)",
                      "A band remains CONSTANT (myosin length does not change)",
                      "Z-lines move CLOSER together",
                      "Key proteins: Actin, Myosin, Tropomyosin, Troponin (TnC, TnI, TnT)",
                      "Energy: ATP needed for cross-bridge cycling AND detachment",
                      "Rigor mortis: No ATP = permanent cross-bridges = stiffness"],
                     color="EBF5FB", border_color="2E86C1")

    # Muscle Fiber Types
    doc.add_paragraph()
    add_heading_styled(doc, "3.5 Types of Skeletal Muscle Fibers", 2)

    create_styled_table(doc,
        ["Feature", "Type I (Slow Oxidative/Red)", "Type IIa (Fast Oxidative)", "Type IIb (Fast Glycolytic/White)"],
        [
            ["Color", "Red (high myoglobin)", "Red-Pink", "White (low myoglobin)"],
            ["Contraction speed", "Slow", "Fast", "Fastest"],
            ["Fatigue resistance", "High (aerobic)", "Moderate", "Low (fatigues quickly)"],
            ["Mitochondria", "Many", "Many", "Few"],
            ["Metabolism", "Oxidative (aerobic)", "Oxidative-Glycolytic", "Glycolytic (anaerobic)"],
            ["Capillary density", "High", "High", "Low"],
            ["Fiber diameter", "Small", "Intermediate", "Large"],
            ["Motor unit size", "Small", "Intermediate", "Large"],
            ["Force production", "Low", "Intermediate", "High"],
            ["Example", "Soleus (postural)", "Mixed muscles", "Extraocular, Gastrocnemius"],
            ["Activity type", "Marathon, Posture", "Swimming, Walking", "Sprinting, Weight lifting"],
        ],
        header_color="1E8449")

    doc.add_page_break()

    # Cardiac Muscle
    add_heading_styled(doc, "3.6 Cardiac Muscle - Special Features", 2)

    add_bullet_points(doc, [
        "Autorhythmicity: Generates own impulses; SA node is natural pacemaker (72 bpm)",
        "Intercalated discs: Unique to cardiac muscle; contain gap junctions for rapid impulse conduction",
        "Functional syncytium: Acts as single unit due to gap junctions (all-or-none contraction)",
        "Long refractory period: Prevents tetanus (sustained contraction); protects heart function",
        "Aerobic metabolism: Almost exclusively oxidative; abundant mitochondria (25-35% of cell volume)",
        "Frank-Starling Law: Greater stretch (preload) = Greater force of contraction",
        "No motor end plates: Modified by ANS but not initiated by it",
        "Calcium-induced calcium release (CICR): Ca2+ entry triggers more Ca2+ release from SR"
    ], bold_prefix=True)

    doc.add_paragraph()
    add_highlight_box(doc, "CLINICAL: Cardiac Muscle Pathology",
                     ["Myocardial infarction (MI): Death of cardiac muscle due to ischemia; replaced by scar tissue",
                      "Cardiac hypertrophy: Enlargement of existing fibers (not hyperplasia)",
                      "Cardiomyopathy: Disease of heart muscle (dilated, hypertrophic, restrictive types)",
                      "Troponin I and T: Cardiac-specific biomarkers for MI diagnosis",
                      "Cardiac muscle cannot regenerate - damage is permanent (fibrosis)"],
                     color="F9EBEA", border_color="CB4335")

    # Smooth Muscle
    doc.add_paragraph()
    add_heading_styled(doc, "3.7 Smooth Muscle - Special Features", 2)

    add_bullet_points(doc, [
        "No sarcomeres: Contractile proteins in dense bodies (equivalent to Z-lines)",
        "Calmodulin: Replaces troponin as calcium-binding regulatory protein",
        "Mechanism: Ca2+-Calmodulin activates MLCK (Myosin Light Chain Kinase) which phosphorylates myosin",
        "Latch mechanism: Maintain tension with minimal ATP expenditure (energy efficient)",
        "Plasticity (stress-relaxation): Can accommodate volume changes (urinary bladder)",
        "Types: Single-unit (visceral - gap junctions, acts as syncytium) & Multi-unit (independent, no gap junctions)",
        "Single-unit: GIT, uterus, ureter, urinary bladder - spontaneous rhythmic activity",
        "Multi-unit: Iris, ciliary body, vas deferens, arrector pili, large airways - fine neural control",
        "Regeneration: Good capacity; can divide (hyperplasia in pregnant uterus)"
    ], bold_prefix=True)

    doc.add_paragraph()
    add_highlight_box(doc, "HIGH YIELD: Smooth vs Skeletal Muscle Contraction",
                     ["Skeletal: Ca2+ binds Troponin C -> exposes actin binding site -> cross-bridge cycling",
                      "Smooth: Ca2+ binds Calmodulin -> activates MLCK -> phosphorylates myosin -> cross-bridge cycling",
                      "Smooth muscle: SLOWER but more SUSTAINED contraction (latch state)",
                      "Smooth muscle can contract to greater degree (shorten more than skeletal)"],
                     color="F4ECF7", border_color="8E44AD")

    doc.add_page_break()


def add_nervous_tissue(doc):
    """Add comprehensive nervous tissue section."""
    add_heading_styled(doc, "SECTION 4: NERVOUS TISSUE", 1)
    add_section_divider(doc)

    add_body_text(doc, "Nervous tissue is the most complex tissue in the body. It is specialized for "
                  "rapid communication via electrical and chemical signals. It forms the brain, spinal cord, "
                  "and peripheral nerves. It is derived from the neuroectoderm (ectoderm).")

    # General Features
    add_heading_styled(doc, "4.1 General Features", 2)
    add_bullet_points(doc, [
        "Two cell types: Neurons (nerve cells) + Neuroglia (supporting cells)",
        "Properties: Excitability and Conductivity (unique to nervous tissue)",
        "Origin: Neuroectoderm (neurons & CNS glia), Neural crest (PNS glia, Schwann cells)",
        "Organization: CNS (brain + spinal cord) & PNS (cranial + spinal nerves, ganglia)",
        "Neurons are post-mitotic: Cannot divide in adults (very limited neurogenesis in hippocampus & SVZ)",
        "Neuroglia can divide: Retain mitotic ability (source of brain tumors - gliomas)"
    ], bold_prefix=True)

    # Neurons
    doc.add_paragraph()
    add_heading_styled(doc, "4.2 Neurons (Nerve Cells)", 2)

    add_body_text(doc, "Neurons are the structural and functional units of the nervous system. "
                  "They generate and transmit nerve impulses (action potentials).")

    add_heading_styled(doc, "Structure of a Typical Neuron:", 3)

    create_styled_table(doc,
        ["Part", "Structure", "Function"],
        [
            ["Cell Body\n(Soma/Perikaryon)", "Contains nucleus (large, pale, central nucleolus)\nNissl bodies (rough ER + ribosomes)\nNeurofibrils, Golgi, Mitochondria\nLipofuscin (wear & tear pigment)",
             "Metabolic center of neuron\nProtein synthesis\nIntegration of signals\nContains most organelles"],
            ["Dendrites", "Short, branching processes\nContain Nissl bodies\nDendritic spines (synaptic sites)\nMany per neuron (usually)",
             "RECEIVE impulses\n(toward cell body)\nIncrease surface area\nfor receiving signals"],
            ["Axon\n(Nerve fiber)", "Single, long process\nArises from axon hillock\nNO Nissl bodies in axon\nAxon hillock: trigger zone\nAxon terminals (synaptic knobs/boutons)",
             "CONDUCT impulses\n(away from cell body)\nAxonal transport\n(anterograde & retrograde)"],
            ["Myelin Sheath", "Lipid-rich insulating layer\nFormed by Schwann cells (PNS)\nOligodendrocytes (CNS)\nNodes of Ranvier: gaps",
             "Increases conduction speed\nSaltatory conduction\nElectrical insulation\nProtection"],
            ["Synaptic Terminals", "Contain synaptic vesicles\nNeurotransmitters stored\nMitochondria (ATP for NT release)\nActive zones",
             "Release neurotransmitters\nSignal transmission to\nnext neuron/effector"],
        ],
        header_color="1B4F72")

    doc.add_paragraph()
    add_highlight_box(doc, "HIGH YIELD: Nissl Bodies (Tigroid Bodies)",
                     ["Basophilic granules in cell body and dendrites (NOT in axon or axon hillock)",
                      "Composed of: Rough ER + Polyribosomes (active protein synthesis)",
                      "Named after Franz Nissl (1860-1919)",
                      "Stains with: Basic dyes (cresyl violet, methylene blue, toluidine blue)",
                      "Chromatolysis: Dissolution of Nissl bodies after axonal injury (axonal reaction)",
                      "  - Occurs 24-48 hours after injury",
                      "  - Nucleus moves to periphery (eccentric)",
                      "  - Cell body swells",
                      "  - Indicates active protein synthesis for repair"],
                     color="E8F8F5", border_color="1ABC9C")

    doc.add_page_break()


    # Classification of Neurons
    add_heading_styled(doc, "4.3 Classification of Neurons", 2)

    add_heading_styled(doc, "A. Structural Classification (Based on Number of Processes):", 3)

    create_styled_table(doc,
        ["Type", "Structure", "Location", "Example"],
        [
            ["Unipolar (Pseudounipolar)", "Single process that divides\ninto peripheral & central branch\nNo true dendrite from soma",
             "Dorsal root ganglia (DRG)\nCranial nerve ganglia\nSensory neurons",
             "Sensory neurons carrying\npain, touch, temperature\nfrom body to spinal cord"],
            ["Bipolar", "One axon + One dendrite\nFrom opposite poles of cell body\nOval cell body",
             "Retina (bipolar cells)\nCochlear & Vestibular ganglia\nOlfactory epithelium",
             "Special sensory neurons\n(vision, hearing,\nsmell, balance)"],
            ["Multipolar", "One axon + Multiple dendrites\nMost common neuron type\n(99% of all neurons)",
             "CNS (brain & spinal cord)\nMotor neurons\nInterneurons\nAutonomic ganglia",
             "Motor neurons\nPurkinje cells (cerebellum)\nPyramidal cells (cortex)\nAnterior horn cells"],
        ],
        header_color="6C3483")

    doc.add_paragraph()
    add_heading_styled(doc, "B. Functional Classification:", 3)

    create_styled_table(doc,
        ["Type", "Function", "Direction of Impulse", "Location"],
        [
            ["Sensory (Afferent)", "Carry impulses from\nreceptors to CNS",
             "Receptor -> CNS\n(periphery to center)", "Dorsal root ganglia\nCranial nerve ganglia\nMost are pseudounipolar"],
            ["Motor (Efferent)", "Carry impulses from\nCNS to effectors",
             "CNS -> Effector\n(center to periphery)", "Anterior horn (somatic)\nLateral horn (autonomic)\nAll are multipolar"],
            ["Interneurons\n(Association/Relay)", "Connect sensory &\nmotor neurons\nIntegration & processing",
             "Within CNS only\n(neuron to neuron)", "Brain & Spinal cord\n99% of all neurons\nAll are multipolar"],
        ],
        header_color="1E8449")

    # Neuroglia
    doc.add_paragraph()
    add_heading_styled(doc, "4.4 Neuroglia (Glial Cells)", 2)

    add_body_text(doc, "Neuroglia ('nerve glue') are non-neuronal cells that support, protect, and maintain "
                  "neurons. They outnumber neurons approximately 10:1 and retain the ability to divide.")

    add_heading_styled(doc, "CNS Glial Cells:", 3)

    create_styled_table(doc,
        ["Cell Type", "Origin", "Function", "Special Features"],
        [
            ["Astrocytes\n(Largest glia)", "Neuroectoderm", "Structural support\nBlood-Brain Barrier (BBB)\nNeurotransmitter uptake\nIon homeostasis (K+ buffering)\nScar formation (gliosis)\nMetabolic support to neurons",
             "Star-shaped with processes\nGFAP marker (specific)\nProtoplasmic (gray matter)\nFibrous (white matter)\nPerivascular feet on blood vessels\nForm BBB with endothelium"],
            ["Oligodendrocytes", "Neuroectoderm", "Form myelin in CNS\nOne cell myelinates\nmultiple axons (up to 50)\nMaintain CNS myelin",
             "Fewer processes than astrocytes\nEquivalent of Schwann cell in PNS\nBUT one oligodendrocyte: many axons\nvs. One Schwann cell: one internode"],
            ["Microglia\n(Smallest glia)", "Mesoderm\n(monocyte lineage)", "Phagocytosis\nImmune defense in CNS\nAntigen presentation\nInflammatory response",
             "ONLY non-neuroectodermal glia\nFrom bone marrow monocytes\nActivated in disease/injury\nHIV reservoir in brain\nRod cells in chronic disease"],
            ["Ependymal Cells", "Neuroectoderm", "Line ventricles & central canal\nCSF circulation (cilia)\nForm choroid plexus\n(CSF production)",
             "Simple columnar/cuboidal\nCiliated (move CSF)\nTanycytes (modified type)\nNo basement membrane\n(allows CSF-brain exchange)"],
        ],
        header_color="1A5276")

    doc.add_paragraph()
    add_heading_styled(doc, "PNS Glial Cells:", 3)

    create_styled_table(doc,
        ["Cell Type", "Origin", "Function", "Special Features"],
        [
            ["Schwann Cells\n(Neurolemmocytes)", "Neural crest", "Form myelin in PNS\nOne cell: one internode\nAid in nerve regeneration\nForm neurilemma (outer sheath)",
             "Essential for PNS regeneration\nGuide regenerating axons\nBands of Bungner (after injury)\nNodes of Ranvier between cells"],
            ["Satellite Cells", "Neural crest", "Surround neuron cell bodies\nin ganglia\nSupport & nourishment\nIon homeostasis",
             "Found in DRG & autonomic ganglia\nAnalogous to astrocytes of CNS\nFlat cells surrounding soma\nRegulate microenvironment"],
        ],
        header_color="7D3C98")

    doc.add_paragraph()
    add_highlight_box(doc, "CLINICAL: Demyelinating Diseases",
                     ["Multiple Sclerosis (MS): Autoimmune demyelination in CNS (oligodendrocytes destroyed)",
                      "  - Relapsing-remitting course; optic neuritis, limb weakness",
                      "Guillain-Barre Syndrome (GBS): Autoimmune demyelination in PNS (Schwann cells attacked)",
                      "  - Ascending paralysis; post-infectious (often after Campylobacter infection)",
                      "Charcot-Marie-Tooth: Hereditary peripheral neuropathy (Schwann cell/myelin defect)",
                      "Gliomas: Tumors from glial cells (astrocytoma, oligodendroglioma, ependymoma, GBM)"],
                     color="FDEDEC", border_color="E74C3C")

    doc.add_page_break()


    # Synapses
    add_heading_styled(doc, "4.5 Synapses", 2)

    add_body_text(doc, "A synapse is the junction between two neurons or between a neuron and an effector cell "
                  "(muscle or gland). Signals are transmitted across synapses by neurotransmitters.")

    add_heading_styled(doc, "Types of Synapses:", 3)
    add_bullet_points(doc, [
        "Chemical synapse: Neurotransmitter-mediated; unidirectional; synaptic delay (0.5ms); most common in human body",
        "Electrical synapse: Gap junctions; bidirectional; no delay; found in cardiac muscle, smooth muscle, some CNS neurons",
        "Axodendritic: Axon terminal to dendrite (most common type)",
        "Axosomatic: Axon terminal to cell body",
        "Axoaxonic: Axon terminal to another axon (presynaptic inhibition/facilitation)"
    ], bold_prefix=True)

    doc.add_paragraph()
    add_heading_styled(doc, "Neurotransmitters:", 3)

    create_styled_table(doc,
        ["Category", "Neurotransmitter", "Location/Function"],
        [
            ["Amino Acids", "Glutamate", "Major EXCITATORY NT in CNS"],
            ["", "GABA", "Major INHIBITORY NT in brain"],
            ["", "Glycine", "Major INHIBITORY NT in spinal cord"],
            ["Monoamines\n(Biogenic Amines)", "Acetylcholine (ACh)", "NMJ, Parasympathetic, ANS ganglia, CNS (memory)"],
            ["", "Norepinephrine (NE)", "Sympathetic postganglionic, Locus coeruleus (arousal)"],
            ["", "Dopamine", "Substantia nigra (movement), VTA (reward/pleasure)"],
            ["", "Serotonin (5-HT)", "Raphe nuclei (mood, sleep, appetite)"],
            ["", "Histamine", "Hypothalamus (wakefulness, gastric acid)"],
            ["Neuropeptides", "Endorphins/Enkephalins", "Pain modulation (natural opioids)"],
            ["", "Substance P", "Pain transmission in spinal cord"],
            ["Gases", "Nitric Oxide (NO)", "Vasodilation; retrograde messenger"],
            ["Purines", "ATP, Adenosine", "Pain signaling; sleep promotion"],
        ],
        header_color="2C3E50")

    # Nerve Fiber Classification
    doc.add_paragraph()
    add_heading_styled(doc, "4.6 Classification of Nerve Fibers", 2)

    create_styled_table(doc,
        ["Type", "Myelination", "Diameter", "Velocity", "Function"],
        [
            ["A-alpha", "Heavily myelinated", "12-20 um", "70-120 m/s", "Motor (skeletal muscle)\nProprioception"],
            ["A-beta", "Myelinated", "5-12 um", "30-70 m/s", "Touch, Pressure"],
            ["A-gamma", "Myelinated", "3-6 um", "15-30 m/s", "Muscle spindle (motor)"],
            ["A-delta", "Lightly myelinated", "2-5 um", "12-30 m/s", "Sharp/fast pain\nTemperature, Touch"],
            ["B fibers", "Lightly myelinated", "1-3 um", "3-15 m/s", "Preganglionic autonomic"],
            ["C fibers", "Unmyelinated", "0.4-1.2 um", "0.5-2 m/s", "Slow/dull pain\nPostganglionic sympathetic\nTemperature"],
        ],
        header_color="154360")

    doc.add_paragraph()
    add_highlight_box(doc, "HIGH YIELD: Nerve Degeneration & Regeneration",
                     ["Wallerian degeneration: Degeneration of axon DISTAL to injury site",
                      "  - Axon + myelin break down distal to cut",
                      "  - Schwann cells survive & form Bands of Bungner (guide tubes)",
                      "  - Macrophages clear debris",
                      "Chromatolysis: Changes in cell body after injury (Nissl body dissolution, nucleus eccentric)",
                      "Regeneration: Only in PNS (Schwann cells guide); rate ~1-3 mm/day",
                      "CNS regeneration: Very limited (inhibitory factors: Nogo, MAG; no Schwann cells)"],
                     color="EBF5FB", border_color="2E86C1")

    doc.add_page_break()


def add_tissue_repair(doc):
    """Add tissue repair and regeneration section."""
    add_heading_styled(doc, "SECTION 5: TISSUE REPAIR & REGENERATION", 1)
    add_section_divider(doc)

    add_heading_styled(doc, "5.1 Cell Populations Based on Regenerative Capacity", 2)

    create_styled_table(doc,
        ["Category", "Description", "Examples", "Regeneration"],
        [
            ["Labile Cells\n(Continuously dividing)", "Cells that divide throughout life\nHigh mitotic rate\nShort cell cycle",
             "Epithelial cells (skin, GIT)\nBone marrow cells\nLining of urinary tract\nUterine epithelium",
             "Excellent regeneration\nReplace by proliferation\nfrom stem cells"],
            ["Stable Cells\n(Quiescent - G0)", "Normally in G0 phase\nCan re-enter cell cycle\nwhen stimulated",
             "Hepatocytes (liver)\nRenal tubular cells\nFibroblasts\nSmooth muscle\nOsteoblasts, Chondroblasts\nEndothelial cells",
             "Good regeneration\nwhen stimulated\n(e.g., liver regeneration\nafter hepatectomy)"],
            ["Permanent Cells\n(Non-dividing)", "Cannot divide (post-mitotic)\nLeft cell cycle permanently\nReplaced by scar tissue",
             "Neurons (CNS)\nCardiac muscle cells\nSkeletal muscle cells\n(limited satellite cells)",
             "No/minimal regeneration\nDamage = permanent loss\nReplaced by fibrosis\n(scar tissue)"],
        ],
        header_color="1E8449")

    doc.add_paragraph()
    add_heading_styled(doc, "5.2 Types of Tissue Repair", 2)

    add_heading_styled(doc, "A. Regeneration:", 3)
    add_bullet_points(doc, [
        "Replacement of damaged tissue by same type of cells",
        "Requires intact basement membrane/tissue framework",
        "Complete restoration of structure and function",
        "Best in labile and stable cells; minimal in permanent cells"
    ])

    add_heading_styled(doc, "B. Fibrosis (Scarring):", 3)
    add_bullet_points(doc, [
        "Replacement by connective tissue (scar)",
        "Occurs when framework is destroyed or permanent cells are damaged",
        "Does NOT restore original function",
        "Fibroblasts deposit collagen (mainly Type I and Type III)"
    ])

    doc.add_paragraph()
    add_heading_styled(doc, "5.3 Wound Healing", 2)

    add_heading_styled(doc, "Phases of Wound Healing:", 3)
    add_numbered_list(doc, [
        "Hemostasis (minutes): Vasoconstriction, platelet plug, fibrin clot",
        "Inflammation (1-3 days): Neutrophils (first 24h), then Macrophages (48-96h); debris removal",
        "Proliferation (3-21 days): Granulation tissue formation, angiogenesis, fibroblast proliferation, epithelialization",
        "Remodeling/Maturation (21 days - 1 year): Collagen reorganization, Type III replaced by Type I, wound contraction, scar maturation"
    ])

    doc.add_paragraph()
    add_highlight_box(doc, "HIGH YIELD: Healing Types",
                     ["Primary intention (First intention): Clean, sutured wound; minimal scar; edges approximated",
                      "Secondary intention (Second intention): Open wound; heals from base up; granulation tissue; larger scar",
                      "Tertiary intention (Delayed primary): Wound left open initially, then closed later",
                      "Granulation tissue: New capillaries + fibroblasts + inflammatory cells (pink, granular appearance)",
                      "Keloid: Excessive collagen deposition BEYOND wound margins (common in dark-skinned individuals)",
                      "Hypertrophic scar: Excessive collagen WITHIN wound margins (does not extend beyond)"],
                     color="FEF9E7", border_color="F39C12")

    doc.add_page_break()


def add_membranes_section(doc):
    """Add body membranes section."""
    add_heading_styled(doc, "SECTION 6: BODY MEMBRANES", 1)
    add_section_divider(doc)

    add_body_text(doc, "Body membranes are thin sheets of tissue that cover body surfaces, line body cavities, "
                  "and surround organs. They are composed of epithelium + underlying connective tissue.")

    create_styled_table(doc,
        ["Membrane Type", "Composition", "Location", "Function"],
        [
            ["Mucous Membrane\n(Mucosa)", "Epithelium + Lamina propria\n(loose CT)\n+ Muscularis mucosae (GIT)",
             "Lines all body cavities that\nopen to exterior:\nGIT, Respiratory, Urinary,\nReproductive tracts",
             "Protection\nSecretion (mucus)\nAbsorption"],
            ["Serous Membrane\n(Serosa)", "Simple squamous epithelium\n(mesothelium) + thin CT layer\nSecretes serous fluid",
             "Lines CLOSED body cavities:\nPleura (lungs)\nPericardium (heart)\nPeritoneum (abdomen)",
             "Reduce friction\nParietal layer (cavity wall)\nVisceral layer (organ surface)\nSerous fluid between layers"],
            ["Cutaneous Membrane\n(Skin)", "Keratinized stratified squamous\nepithelium (epidermis) +\nDermis (CT)",
             "Covers entire external\nbody surface",
             "Protection\nTemperature regulation\nSensation\nVitamin D synthesis"],
            ["Synovial Membrane", "NO epithelial lining\nSpecialized CT only\n(Type A & B synoviocytes)",
             "Lines joint cavities\nBursae\nTendon sheaths",
             "Produce synovial fluid\n(lubricant for joints)\nNutrish avascular cartilage"],
        ],
        header_color="7D3C98")

    doc.add_paragraph()
    add_highlight_box(doc, "REMEMBER: Serous Membrane Naming",
                     ["Pleura: Parietal pleura (chest wall) + Visceral pleura (lung surface); Pleural cavity between",
                      "Pericardium: Parietal pericardium (fibrous sac) + Visceral pericardium (epicardium); Pericardial cavity",
                      "Peritoneum: Parietal peritoneum (abdominal wall) + Visceral peritoneum (organs); Peritoneal cavity",
                      "Mesentery: Double layer of peritoneum suspending organs",
                      "Clinical: Pleuritis, Pericarditis, Peritonitis = inflammation of serous membranes"],
                     color="F4ECF7", border_color="8E44AD")

    doc.add_page_break()


def add_mcq_section(doc):
    """Add previous year MCQs section."""
    add_heading_styled(doc, "SECTION 7: PREVIOUS YEAR MCQs", 1)
    add_section_divider(doc)

    add_body_text(doc, "The following MCQs have been asked in various competitive examinations including "
                  "NEET, AIIMS, JIPMER, GPAT, UPPSC, APPSC, and University exams. Practice these for exam preparation.")

    doc.add_paragraph()
    add_heading_styled(doc, "7.1 Epithelial Tissue MCQs", 2)

    mcqs_epithelial = [
        ("Q1. Transitional epithelium is found in: [NEET 2019]",
         "(a) Trachea  (b) Urinary bladder  (c) Stomach  (d) Small intestine",
         "Answer: (b) Urinary bladder\nExplanation: Transitional epithelium (urothelium) lines the urinary tract - renal pelvis, ureter, urinary bladder, and upper urethra. It can stretch without rupturing."),

        ("Q2. Pseudostratified ciliated columnar epithelium lines: [AIIMS 2018]",
         "(a) Stomach  (b) Trachea  (c) Small intestine  (d) Esophagus",
         "Answer: (b) Trachea\nExplanation: The trachea and primary bronchi are lined by pseudostratified ciliated columnar epithelium with goblet cells. This forms the mucociliary escalator."),

        ("Q3. Goblet cells secrete: [GPAT 2020]",
         "(a) Serous fluid  (b) Mucus  (c) Sebum  (d) Sweat",
         "Answer: (b) Mucus\nExplanation: Goblet cells are unicellular exocrine glands that secrete mucin (glycoprotein), which when hydrated forms mucus. Found in respiratory and GI tract."),

        ("Q4. Which epithelium lines blood vessels? [NEET 2017]",
         "(a) Simple squamous  (b) Simple cuboidal  (c) Stratified squamous  (d) Pseudostratified",
         "Answer: (a) Simple squamous (Endothelium)\nExplanation: The endothelium is simple squamous epithelium lining all blood vessels, lymphatics, and the heart."),

        ("Q5. Kartagener syndrome is due to defect in: [AIIMS 2020]",
         "(a) Microvilli  (b) Cilia (dynein arms)  (c) Stereocilia  (d) Flagella",
         "Answer: (b) Cilia (dynein arms)\nExplanation: Kartagener syndrome (Primary Ciliary Dyskinesia) is due to absent/defective dynein arms in cilia. Triad: Situs inversus + Bronchiectasis + Sinusitis."),

        ("Q6. Stratified squamous non-keratinized epithelium is found in: [JIPMER 2019]",
         "(a) Skin  (b) Esophagus  (c) Stomach  (d) Trachea",
         "Answer: (b) Esophagus\nExplanation: Esophagus is lined by stratified squamous non-keratinized epithelium (moist, no keratin) for protection against abrasion during swallowing."),

        ("Q7. Holocrine glands include: [GPAT 2019]",
         "(a) Salivary glands  (b) Sweat glands  (c) Sebaceous glands  (d) Mammary glands",
         "Answer: (c) Sebaceous glands\nExplanation: In holocrine secretion, the entire cell ruptures and dies to release its secretion. Sebaceous glands are the classic example."),

        ("Q8. The basement membrane is composed of: [NEET 2018]",
         "(a) Type I collagen  (b) Type IV collagen  (c) Elastin  (d) Type III collagen",
         "Answer: (b) Type IV collagen\nExplanation: Basement membrane (basal lamina) contains Type IV collagen, laminin, fibronectin, and heparan sulfate proteoglycan (perlecan)."),

        ("Q9. Tight junctions (zonula occludens) contain which protein? [AIIMS 2017]",
         "(a) Cadherins  (b) Integrins  (c) Claudins and Occludins  (d) Connexins",
         "Answer: (c) Claudins and Occludins\nExplanation: Tight junctions are formed by claudins and occludins that create an impermeable seal between cells, preventing paracellular transport."),

        ("Q10. Pemphigus vulgaris involves antibodies against: [NEET 2020]",
         "(a) Hemidesmosomes  (b) Desmosomes (Desmoglein)  (c) Tight junctions  (d) Gap junctions",
         "Answer: (b) Desmosomes (Desmoglein 3)\nExplanation: Pemphigus vulgaris has IgG autoantibodies against desmoglein 3 (in desmosomes), causing acantholysis (loss of cell-cell adhesion) and intraepidermal blisters."),
    ]

    for q, options, answer in mcqs_epithelial:
        p = doc.add_paragraph()
        run = p.add_run(q)
        run.bold = True
        run.font.size = Pt(10)

        p = doc.add_paragraph()
        run = p.add_run(options)
        run.font.size = Pt(10)

        p = doc.add_paragraph()
        run = p.add_run(answer)
        run.font.size = Pt(9.5)
        run.italic = True
        run.font.color.rgb = RGBColor(0x1E, 0x84, 0x49)

        doc.add_paragraph()

    doc.add_page_break()


    # Connective Tissue MCQs
    add_heading_styled(doc, "7.2 Connective Tissue MCQs", 2)

    mcqs_connective = [
        ("Q11. Most abundant protein in human body is: [NEET 2019]",
         "(a) Actin  (b) Myosin  (c) Collagen  (d) Albumin",
         "Answer: (c) Collagen\nExplanation: Collagen constitutes 25-30% of total body protein, making it the most abundant protein. Type I collagen is the most common type."),

        ("Q12. Mast cells release: [AIIMS 2018]",
         "(a) Histamine & Heparin  (b) Antibodies  (c) Collagen  (d) Melanin",
         "Answer: (a) Histamine & Heparin\nExplanation: Mast cells contain metachromatic granules with histamine (vasodilator), heparin (anticoagulant), and proteases. They mediate Type I hypersensitivity (anaphylaxis)."),

        ("Q13. Scurvy is caused by deficiency of: [GPAT 2020]",
         "(a) Vitamin A  (b) Vitamin C  (c) Vitamin D  (d) Vitamin K",
         "Answer: (b) Vitamin C\nExplanation: Vitamin C is essential for hydroxylation of proline and lysine in collagen synthesis. Deficiency causes scurvy (bleeding gums, poor wound healing, weakened CT)."),

        ("Q14. Marfan syndrome is due to defect in: [NEET 2020]",
         "(a) Type I collagen  (b) Type IV collagen  (c) Fibrillin-1  (d) Elastin",
         "Answer: (c) Fibrillin-1\nExplanation: Marfan syndrome is autosomal dominant with defect in FBN1 gene (Fibrillin-1, component of elastic fibers). Features: tall stature, arachnodactyly, lens subluxation, aortic aneurysm."),

        ("Q15. Which cartilage never calcifies? [JIPMER 2018]",
         "(a) Hyaline  (b) Elastic  (c) Fibrocartilage  (d) All can calcify",
         "Answer: (b) Elastic cartilage\nExplanation: Elastic cartilage NEVER undergoes calcification. Hyaline cartilage commonly calcifies (e.g., costal cartilage in elderly, growth plate closure). Fibrocartilage may calcify in pathology."),

        ("Q16. Clock-face (cartwheel) nucleus is seen in: [AIIMS 2019]",
         "(a) Plasma cells  (b) Mast cells  (c) Macrophages  (d) Eosinophils",
         "Answer: (a) Plasma cells\nExplanation: Plasma cells have characteristic clock-face/cartwheel nucleus (heterochromatin arranged like spokes of wheel), abundant RER (basophilic), and perinuclear halo (Golgi area)."),

        ("Q17. Wharton's jelly is an example of: [GPAT 2019]",
         "(a) Areolar tissue  (b) Mucous connective tissue  (c) Adipose tissue  (d) Reticular tissue",
         "Answer: (b) Mucous connective tissue\nExplanation: Wharton's jelly is mucous CT found in umbilical cord. It is a type of embryonic CT with fibroblasts in abundant gelatinous ground substance rich in hyaluronic acid."),

        ("Q18. Brown fat is characterized by: [NEET 2018]",
         "(a) Unilocular fat droplets  (b) Multilocular fat droplets  (c) No mitochondria  (d) No blood supply",
         "Answer: (b) Multilocular fat droplets\nExplanation: Brown adipose tissue (BAT) has multilocular fat droplets, abundant mitochondria (give brown color due to cytochromes), rich blood supply, and UCP-1 (thermogenin) for non-shivering thermogenesis."),

        ("Q19. Achondroplasia is caused by mutation in: [AIIMS 2020]",
         "(a) FGFR3  (b) FGFR1  (c) Type II collagen  (d) Aggrecan",
         "Answer: (a) FGFR3 (Fibroblast Growth Factor Receptor 3)\nExplanation: Achondroplasia (most common cause of dwarfism) is due to gain-of-function mutation in FGFR3 gene. It inhibits endochondral ossification at growth plates."),

        ("Q20. Osteons (Haversian systems) are found in: [JIPMER 2020]",
         "(a) Compact bone  (b) Spongy bone  (c) Cartilage  (d) Both a and b",
         "Answer: (a) Compact bone\nExplanation: Osteons (Haversian systems) are the structural units of compact bone. They consist of concentric lamellae of bone around a central Haversian canal containing blood vessels."),
    ]

    for q, options, answer in mcqs_connective:
        p = doc.add_paragraph()
        run = p.add_run(q)
        run.bold = True
        run.font.size = Pt(10)

        p = doc.add_paragraph()
        run = p.add_run(options)
        run.font.size = Pt(10)

        p = doc.add_paragraph()
        run = p.add_run(answer)
        run.font.size = Pt(9.5)
        run.italic = True
        run.font.color.rgb = RGBColor(0x1E, 0x84, 0x49)

        doc.add_paragraph()

    doc.add_page_break()


    # Muscular Tissue MCQs
    add_heading_styled(doc, "7.3 Muscular Tissue MCQs", 2)

    mcqs_muscular = [
        ("Q21. Intercalated discs are found in: [NEET 2019]",
         "(a) Skeletal muscle  (b) Cardiac muscle  (c) Smooth muscle  (d) All muscles",
         "Answer: (b) Cardiac muscle\nExplanation: Intercalated discs are unique to cardiac muscle. They contain desmosomes (adhesion), gap junctions (electrical coupling), and fascia adherens (mechanical coupling)."),

        ("Q22. The functional unit of muscle contraction is: [AIIMS 2018]",
         "(a) Myofibril  (b) Sarcomere  (c) Muscle fiber  (d) Fascicle",
         "Answer: (b) Sarcomere\nExplanation: The sarcomere (Z-line to Z-line) is the basic contractile unit. During contraction, I band and H zone shorten while A band remains constant."),

        ("Q23. Which band remains constant during muscle contraction? [NEET 2020]",
         "(a) I band  (b) A band  (c) H zone  (d) All change",
         "Answer: (b) A band\nExplanation: A band (dark band) contains thick myosin filaments whose length does not change. I band and H zone shorten as thin filaments slide over thick filaments."),

        ("Q24. Smooth muscle contraction is regulated by: [GPAT 2019]",
         "(a) Troponin  (b) Calmodulin  (c) Tropomyosin only  (d) Titin",
         "Answer: (b) Calmodulin\nExplanation: In smooth muscle, Ca2+ binds to calmodulin (not troponin - which is absent). Ca2+-Calmodulin complex activates MLCK (Myosin Light Chain Kinase) which phosphorylates myosin for contraction."),

        ("Q25. Satellite cells in skeletal muscle are responsible for: [AIIMS 2019]",
         "(a) Contraction  (b) Regeneration/Repair  (c) Energy storage  (d) Nerve supply",
         "Answer: (b) Regeneration/Repair\nExplanation: Satellite cells are stem cells located between the sarcolemma and basement membrane of skeletal muscle fibers. They activate for repair and limited regeneration after injury."),

        ("Q26. Rigor mortis occurs because: [NEET 2017]",
         "(a) Excess calcium  (b) Lack of ATP  (c) Excess ATP  (d) Lack of calcium",
         "Answer: (b) Lack of ATP\nExplanation: After death, ATP is depleted. Without ATP, myosin cannot detach from actin (ATP needed for detachment step in cross-bridge cycle), causing permanent rigidity."),

        ("Q27. T-tubules in cardiac muscle are located at: [JIPMER 2019]",
         "(a) A-I junction  (b) Z-line  (c) M-line  (d) H-zone",
         "Answer: (b) Z-line\nExplanation: In cardiac muscle, T-tubules are at Z-lines forming DYADS (1 T-tubule + 1 SR terminal cisterna). In skeletal muscle, T-tubules are at A-I junction forming TRIADS."),

        ("Q28. Frank-Starling law states: [AIIMS 2020]",
         "(a) Greater stretch = Greater force of contraction  (b) All-or-none law  (c) Rate of contraction increases with temperature  (d) Force decreases with stretch",
         "Answer: (a) Greater stretch = Greater force of contraction\nExplanation: Within physiological limits, greater stretch (increased venous return/preload) of cardiac muscle fibers leads to greater force of contraction. This is the basis of the Frank-Starling mechanism."),

        ("Q29. Type I (slow oxidative) muscle fibers are characterized by: [GPAT 2020]",
         "(a) White color, fast twitch  (b) Red color, fatigue resistant  (c) White color, glycolytic  (d) Large diameter, fast contraction",
         "Answer: (b) Red color, fatigue resistant\nExplanation: Type I fibers are red (high myoglobin), slow-twitch, fatigue-resistant, with many mitochondria and rich capillary supply. Suited for endurance activities and postural maintenance."),

        ("Q30. Multinucleated giant cells derived from muscle are called: [NEET 2018]",
         "(a) Osteoclasts  (b) Skeletal muscle fibers  (c) Megakaryocytes  (d) Syncytiotrophoblasts",
         "Answer: (b) Skeletal muscle fibers\nExplanation: Skeletal muscle fibers are multinucleated cells formed by fusion of myoblasts during development. Nuclei are located peripherally beneath the sarcolemma."),
    ]

    for q, options, answer in mcqs_muscular:
        p = doc.add_paragraph()
        run = p.add_run(q)
        run.bold = True
        run.font.size = Pt(10)

        p = doc.add_paragraph()
        run = p.add_run(options)
        run.font.size = Pt(10)

        p = doc.add_paragraph()
        run = p.add_run(answer)
        run.font.size = Pt(9.5)
        run.italic = True
        run.font.color.rgb = RGBColor(0x1E, 0x84, 0x49)

        doc.add_paragraph()

    doc.add_page_break()


    # Nervous Tissue MCQs
    add_heading_styled(doc, "7.4 Nervous Tissue MCQs", 2)

    mcqs_nervous = [
        ("Q31. Nissl bodies are absent in: [NEET 2020]",
         "(a) Cell body  (b) Dendrites  (c) Axon  (d) Both a and b",
         "Answer: (c) Axon\nExplanation: Nissl bodies (rough ER + ribosomes) are present in the cell body and dendrites but ABSENT in the axon and axon hillock. This is a very frequently asked fact."),

        ("Q32. Myelin sheath in CNS is formed by: [AIIMS 2019]",
         "(a) Schwann cells  (b) Oligodendrocytes  (c) Astrocytes  (d) Microglia",
         "Answer: (b) Oligodendrocytes\nExplanation: In CNS, oligodendrocytes form myelin (one cell myelinates multiple axons). In PNS, Schwann cells form myelin (one cell: one internode of one axon)."),

        ("Q33. The only glial cell of mesodermal origin is: [JIPMER 2019]",
         "(a) Astrocyte  (b) Oligodendrocyte  (c) Microglia  (d) Ependymal cell",
         "Answer: (c) Microglia\nExplanation: Microglia are derived from monocyte/mesoderm lineage. All other CNS glial cells (astrocytes, oligodendrocytes, ependymal cells) are derived from neuroectoderm."),

        ("Q34. Blood-Brain Barrier (BBB) is formed by: [NEET 2018]",
         "(a) Microglia  (b) Astrocyte end-feet  (c) Oligodendrocytes  (d) Ependymal cells",
         "Answer: (b) Astrocyte end-feet (perivascular feet)\nExplanation: BBB is formed by tight junctions of endothelial cells INDUCED by astrocyte end-feet. The structural basis is endothelial tight junctions, but astrocytes are essential for their formation and maintenance."),

        ("Q35. Saltatory conduction occurs in: [GPAT 2020]",
         "(a) Unmyelinated fibers  (b) Myelinated fibers  (c) Both  (d) Smooth muscle",
         "Answer: (b) Myelinated fibers\nExplanation: In myelinated nerves, action potentials jump from one Node of Ranvier to the next (saltatory = leaping). This dramatically increases conduction velocity and conserves energy."),

        ("Q36. Gap junctions are composed of: [AIIMS 2018]",
         "(a) Claudins  (b) Connexins  (c) Cadherins  (d) Integrins",
         "Answer: (b) Connexins\nExplanation: Gap junctions consist of connexons (hemichannels). Each connexon is made of 6 connexin proteins. They allow direct cell-to-cell communication for ions and small molecules."),

        ("Q37. Chromatolysis refers to: [NEET 2019]",
         "(a) Dissolution of Nissl bodies  (b) Breakdown of myelin  (c) Death of neurons  (d) Loss of chromosomes",
         "Answer: (a) Dissolution of Nissl bodies\nExplanation: Chromatolysis occurs in the cell body after axonal injury. Nissl bodies dissolve (disperse), nucleus moves peripherally, and cell body swells. It indicates active protein synthesis for repair."),

        ("Q38. Wallerian degeneration occurs: [JIPMER 2020]",
         "(a) Proximal to injury  (b) Distal to injury  (c) At the site of injury  (d) In cell body",
         "Answer: (b) Distal to injury\nExplanation: Wallerian degeneration affects the portion of axon DISTAL to the site of injury. The axon and myelin sheath degenerate, while Schwann cells survive and form guides for regeneration."),

        ("Q39. The major inhibitory neurotransmitter in brain is: [GPAT 2019]",
         "(a) Glutamate  (b) GABA  (c) Acetylcholine  (d) Dopamine",
         "Answer: (b) GABA (Gamma-Aminobutyric Acid)\nExplanation: GABA is the main inhibitory neurotransmitter in the brain. Glycine is the main inhibitory NT in the spinal cord. Glutamate is the main excitatory NT in the CNS."),

        ("Q40. Multiple Sclerosis involves destruction of: [AIIMS 2020]",
         "(a) Schwann cells  (b) Oligodendrocytes  (c) Astrocytes  (d) Neurons",
         "Answer: (b) Oligodendrocytes (CNS myelin)\nExplanation: MS is an autoimmune demyelinating disease of the CNS. T-cells attack oligodendrocytes/myelin sheaths. Guillain-Barre syndrome is the PNS equivalent (attacks Schwann cells)."),
    ]

    for q, options, answer in mcqs_nervous:
        p = doc.add_paragraph()
        run = p.add_run(q)
        run.bold = True
        run.font.size = Pt(10)

        p = doc.add_paragraph()
        run = p.add_run(options)
        run.font.size = Pt(10)

        p = doc.add_paragraph()
        run = p.add_run(answer)
        run.font.size = Pt(9.5)
        run.italic = True
        run.font.color.rgb = RGBColor(0x1E, 0x84, 0x49)

        doc.add_paragraph()

    doc.add_page_break()


    # Mixed/General MCQs
    add_heading_styled(doc, "7.5 Mixed/General Tissue MCQs", 2)

    mcqs_general = [
        ("Q41. Which cells are called 'dust cells'? [NEET 2018]",
         "(a) Osteoclasts  (b) Alveolar macrophages  (c) Mast cells  (d) Kupffer cells",
         "Answer: (b) Alveolar macrophages\nExplanation: Dust cells are macrophages in lung alveoli that phagocytose inhaled particles, bacteria, and debris. Kupffer cells are liver macrophages; Osteoclasts are bone macrophages."),

        ("Q42. Tendons are made of: [GPAT 2020]",
         "(a) Dense regular CT  (b) Dense irregular CT  (c) Loose CT  (d) Elastic CT",
         "Answer: (a) Dense regular connective tissue\nExplanation: Tendons consist of parallel bundles of Type I collagen fibers (dense regular CT). This arrangement provides great tensile strength in one direction (muscle to bone attachment)."),

        ("Q43. Which tissue type is avascular? [JIPMER 2019]",
         "(a) Bone  (b) Cartilage  (c) Blood  (d) Muscle",
         "Answer: (b) Cartilage\nExplanation: Cartilage is avascular (no blood supply), aneural (no nerves), and alymphatic. Nutrients reach chondrocytes by diffusion from perichondrium. This limits cartilage thickness and repair capacity."),

        ("Q44. Intervertebral discs are made of: [NEET 2017]",
         "(a) Hyaline cartilage  (b) Elastic cartilage  (c) Fibrocartilage  (d) Dense irregular CT",
         "Answer: (c) Fibrocartilage\nExplanation: Intervertebral discs consist of fibrocartilage (annulus fibrosus) surrounding a gelatinous nucleus pulposus. Fibrocartilage provides strength with some compressibility."),

        ("Q45. The fastest conducting nerve fibers are: [AIIMS 2019]",
         "(a) C fibers  (b) A-delta fibers  (c) A-alpha fibers  (d) B fibers",
         "Answer: (c) A-alpha fibers\nExplanation: A-alpha fibers are the largest (12-20 um) and most heavily myelinated, conducting at 70-120 m/s. They carry motor impulses to skeletal muscle and proprioception. C fibers are slowest (0.5-2 m/s)."),

        ("Q46. Osteogenesis imperfecta is caused by defect in: [NEET 2019]",
         "(a) Type I collagen  (b) Type II collagen  (c) Type III collagen  (d) Type IV collagen",
         "Answer: (a) Type I collagen\nExplanation: OI (brittle bone disease) is due to defective Type I collagen synthesis. Features: fractures with minimal trauma, blue sclera, hearing loss, dentinogenesis imperfecta."),

        ("Q47. Meissner's corpuscles are located in: [GPAT 2018]",
         "(a) Dermal papillae  (b) Hypodermis  (c) Deep dermis  (d) Epidermis",
         "Answer: (a) Dermal papillae (superficial dermis)\nExplanation: Meissner's corpuscles are encapsulated mechanoreceptors in dermal papillae. They detect light touch and are abundant in fingertips, lips, and palms. Pacinian corpuscles are in deep dermis/hypodermis (pressure/vibration)."),

        ("Q48. Which is NOT a property of muscle tissue? [NEET 2020]",
         "(a) Contractility  (b) Excitability  (c) Secretion  (d) Elasticity",
         "Answer: (c) Secretion\nExplanation: The four properties of muscle tissue are: Excitability, Contractility, Extensibility, and Elasticity. Secretion is a property of glandular epithelium, not muscle tissue."),

        ("Q49. Ehlers-Danlos syndrome involves: [AIIMS 2017]",
         "(a) Defective elastin  (b) Defective collagen  (c) Defective fibrillin  (d) Defective laminin",
         "Answer: (b) Defective collagen\nExplanation: Ehlers-Danlos syndrome is a group of disorders with defective collagen synthesis/processing. Features: hyperextensible skin, hypermobile joints, tissue fragility, easy bruising. Multiple types with different collagen defects."),

        ("Q50. The largest cell in the human body is: [JIPMER 2018]",
         "(a) Neuron  (b) Ovum  (c) Skeletal muscle fiber  (d) Megakaryocyte",
         "Answer: (c) Skeletal muscle fiber (by length, up to 30-40 cm)\nExplanation: By length, skeletal muscle fibers are longest (up to 40 cm in sartorius). By volume/diameter, ovum is largest single cell (120 um). Neurons can also be very long (motor neurons from spinal cord to foot ~1 meter)."),
    ]

    for q, options, answer in mcqs_general:
        p = doc.add_paragraph()
        run = p.add_run(q)
        run.bold = True
        run.font.size = Pt(10)

        p = doc.add_paragraph()
        run = p.add_run(options)
        run.font.size = Pt(10)

        p = doc.add_paragraph()
        run = p.add_run(answer)
        run.font.size = Pt(9.5)
        run.italic = True
        run.font.color.rgb = RGBColor(0x1E, 0x84, 0x49)

        doc.add_paragraph()

    doc.add_page_break()


def add_quick_revision(doc):
    """Add quick revision section with key facts."""
    add_heading_styled(doc, "SECTION 8: QUICK REVISION - KEY FACTS", 1)
    add_section_divider(doc)

    add_heading_styled(doc, "8.1 Important 'First' and 'Largest' Facts", 2)

    create_styled_table(doc,
        ["Fact", "Answer"],
        [
            ["Most abundant tissue in body", "Connective tissue"],
            ["Most abundant protein in body", "Collagen (Type I - most common type)"],
            ["Most common cell in CT proper", "Fibroblast"],
            ["Hardest tissue in body", "Enamel (then Dentin, then Bone)"],
            ["Hardest CT in body", "Bone"],
            ["Largest cell (by length)", "Skeletal muscle fiber (up to 40 cm)"],
            ["Largest cell (by volume)", "Ovum (120 um diameter)"],
            ["Longest cell", "Motor neuron (up to 1 meter)"],
            ["Smallest cell in body", "Sperm (head: 4-5 um)"],
            ["Most widely distributed epithelium", "Stratified squamous"],
            ["Most common type of cartilage", "Hyaline cartilage"],
            ["Strongest cartilage", "Fibrocartilage"],
            ["Cartilage that never calcifies", "Elastic cartilage"],
            ["Only non-ectodermal glial cell", "Microglia (mesodermal)"],
            ["Most common neuron type", "Multipolar"],
            ["Fastest nerve fiber", "A-alpha (70-120 m/s)"],
            ["Slowest nerve fiber", "C fibers (0.5-2 m/s)"],
            ["Most common WBC", "Neutrophil (60-70%)"],
            ["Father of Histology", "Marie Francois Xavier Bichat"],
        ],
        header_color="1B4F72")

    doc.add_paragraph()
    add_heading_styled(doc, "8.2 Important Staining Reactions", 2)

    create_styled_table(doc,
        ["Structure/Tissue", "Stain", "Color"],
        [
            ["Collagen fibers", "Van Gieson", "Red"],
            ["Collagen fibers", "Masson Trichrome", "Blue/Green"],
            ["Elastic fibers", "Orcein", "Brown"],
            ["Elastic fibers", "Verhoeff", "Black"],
            ["Reticular fibers", "Silver stains (Gomori)", "Black (argyrophilic)"],
            ["Mast cell granules", "Toluidine blue", "Metachromatic (purple-red)"],
            ["Nissl bodies", "Cresyl violet / Methylene blue", "Blue-violet"],
            ["Myelin", "Luxol fast blue", "Blue"],
            ["Glycogen / BM", "PAS (Periodic Acid-Schiff)", "Magenta/Pink"],
            ["Fat/Lipid", "Sudan III/IV, Oil Red O", "Red/Orange"],
            ["Muscle (general)", "H&E (Eosin)", "Pink"],
            ["Nuclei", "Hematoxylin", "Blue/Purple"],
            ["Bone matrix", "H&E (Eosin)", "Pink"],
            ["Cartilage matrix", "H&E", "Purple-blue (basophilic)"],
            ["Amyloid", "Congo Red", "Apple-green birefringence"],
            ["Calcium deposits", "Von Kossa", "Black"],
            ["Iron (hemosiderin)", "Prussian blue (Perl's)", "Blue"],
        ],
        header_color="6C3483")

    doc.add_paragraph()
    add_heading_styled(doc, "8.3 Clinical Correlations Summary", 2)

    create_styled_table(doc,
        ["Disease", "Tissue/Structure Affected", "Defect"],
        [
            ["Scurvy", "Collagen", "Vitamin C deficiency (hydroxylation defect)"],
            ["Osteogenesis imperfecta", "Type I Collagen", "Genetic defect in COL1A1/COL1A2"],
            ["Ehlers-Danlos syndrome", "Various Collagens", "Multiple types with collagen defects"],
            ["Marfan syndrome", "Elastic fibers (Fibrillin-1)", "FBN1 gene mutation"],
            ["Alport syndrome", "Type IV Collagen (BM)", "Glomerulonephritis + hearing loss"],
            ["Achondroplasia", "Growth plate cartilage", "FGFR3 gain-of-function mutation"],
            ["Kartagener syndrome", "Cilia (dynein arms)", "Immotile cilia + situs inversus"],
            ["Pemphigus vulgaris", "Desmosomes (Desmoglein 3)", "Autoantibodies - intraepidermal blisters"],
            ["Bullous pemphigoid", "Hemidesmosomes", "Autoantibodies - subepidermal blisters"],
            ["Osteoporosis", "Bone", "Decreased bone mass/density"],
            ["Osteomalacia/Rickets", "Bone mineralization", "Vitamin D deficiency"],
            ["Multiple sclerosis", "CNS myelin (oligodendrocytes)", "Autoimmune demyelination"],
            ["Guillain-Barre", "PNS myelin (Schwann cells)", "Autoimmune demyelination"],
            ["Duchenne muscular dystrophy", "Skeletal muscle (dystrophin)", "X-linked; absence of dystrophin"],
            ["Myasthenia gravis", "NMJ (ACh receptors)", "Autoantibodies against nicotinic receptors"],
        ],
        header_color="CB4335")

    doc.add_paragraph()
    add_highlight_box(doc, "MNEMONICS COMPILATION",
                     ["ECMN: Four tissue types - Epithelial, Connective, Muscular, Nervous",
                      "MAH: Modes of secretion - Merocrine, Apocrine, Holocrine",
                      "Never Let Monkeys Eat Bananas: WBC order (N>L>M>E>B)",
                      "HAI decreases: During contraction H zone, A-I junction, I band decrease",
                      "Some Killers Have Pretty Nice Weapons: Skin layers (Stratum: Corneum, Lucidum, Granulosum, Spinosum, Basale)",
                      "CAST: Cells that divide continuously - Continuously dividing = LABILE cells",
                      "Sally Got Fudge: Ground substance components - GAGs, Glycoproteins, Fluid"],
                     color="EBF5FB", border_color="2E86C1")

    doc.add_page_break()


def add_additional_high_yield(doc):
    """Add additional high-yield topics."""
    add_heading_styled(doc, "SECTION 9: ADDITIONAL HIGH-YIELD TOPICS", 1)
    add_section_divider(doc)

    add_heading_styled(doc, "9.1 Collagen Types - Complete Table", 2)

    create_styled_table(doc,
        ["Type", "Location", "Associated Disease"],
        [
            ["Type I", "Bone, Skin, Tendon, Dentin, Cornea, Scar tissue\n(Most abundant - 90%)", "Osteogenesis imperfecta\nEhlers-Danlos (some types)"],
            ["Type II", "Hyaline & Elastic cartilage, Vitreous humor, Nucleus pulposus", "Achondrogenesis type II\nStickler syndrome"],
            ["Type III", "Reticular fibers, Blood vessels, Uterus, Skin\n(Fetal tissue - 'Reticulin')", "Ehlers-Danlos type IV (vascular)\nBlood vessel rupture"],
            ["Type IV", "Basement membrane (Basal lamina)", "Alport syndrome\nGoodpasture syndrome (antibodies against it)"],
            ["Type V", "Placenta, Cornea, Interstitial tissue (with Type I)", "Ehlers-Danlos classical type"],
            ["Type VII", "Anchoring fibrils (dermo-epidermal junction)", "Epidermolysis bullosa dystrophica"],
            ["Type IX", "Cartilage (with Type II)", "Multiple epiphyseal dysplasia"],
            ["Type X", "Hypertrophic zone of growth plate", "Schmid metaphyseal chondrodysplasia"],
            ["Type XVII", "Hemidesmosomes", "Bullous pemphigoid"],
        ],
        header_color="1E8449")

    doc.add_paragraph()
    add_highlight_box(doc, "MNEMONIC for Collagen Types",
                     ["Type I: 'Bone, Skin, Tendon' = Be Strong Together",
                      "Type II: Cartilage (Two C's = Type II Cartilage)",
                      "Type III: Reticulin (blood vessels, fetal tissue)",
                      "Type IV: Basement Membrane (4 = Floor/Base)",
                      "Collagen synthesis requires: Vitamin C, Iron, Alpha-ketoglutarate, O2, Cu2+",
                      "Steps: Transcription > Translation > Hydroxylation (Vit C) > Glycosylation > Triple helix > Secretion > Cleavage > Cross-linking (Cu2+/Lysyl oxidase)"],
                     color="F4ECF7", border_color="8E44AD")

    doc.add_paragraph()
    add_heading_styled(doc, "9.2 Comparison: Epithelium vs Connective Tissue", 2)

    create_styled_table(doc,
        ["Feature", "Epithelial Tissue", "Connective Tissue"],
        [
            ["Cellularity", "Highly cellular", "Sparse cells, abundant matrix"],
            ["Intercellular matrix", "Minimal", "Abundant (defines the tissue)"],
            ["Vascularity", "Avascular", "Highly vascular (usually)"],
            ["Basement membrane", "Present (cells rest on BM)", "Absent"],
            ["Cell arrangement", "Closely packed, organized", "Scattered in matrix"],
            ["Location", "Surfaces (covering/lining)", "Deep to epithelium; everywhere"],
            ["Origin", "All 3 germ layers", "Mesoderm (mesenchyme)"],
            ["Regeneration", "High (labile cells)", "Variable (depends on type)"],
            ["Cell junctions", "Prominent (tight, desmosomes)", "Few (cells widely separated)"],
            ["Polarity", "Present (apical-basal)", "Absent"],
        ],
        header_color="2C3E50")

    doc.add_paragraph()
    add_heading_styled(doc, "9.3 Stem Cells in Tissues", 2)

    create_styled_table(doc,
        ["Tissue", "Stem Cell", "Location", "Produces"],
        [
            ["Skin (Epidermis)", "Basal cells (stratum basale)", "Deepest layer of epidermis", "Keratinocytes"],
            ["Intestinal epithelium", "Crypt base columnar cells", "Base of crypts of Lieberkuhn", "All intestinal cell types"],
            ["Blood", "Hematopoietic stem cells (HSCs)", "Bone marrow", "All blood cells"],
            ["Bone", "Osteoprogenitor cells", "Periosteum, Endosteum", "Osteoblasts > Osteocytes"],
            ["Cartilage", "Chondrogenic cells", "Perichondrium", "Chondroblasts > Chondrocytes"],
            ["Skeletal muscle", "Satellite cells", "Between sarcolemma & BM", "Myoblasts > Muscle fibers"],
            ["Nervous tissue", "Neural stem cells", "Subventricular zone, Hippocampus", "Neurons, Glia (limited)"],
            ["Liver", "Hepatic oval cells", "Canal of Hering", "Hepatocytes, Bile duct cells"],
            ["Connective tissue", "Mesenchymal stem cells", "Perivascular (pericytes)", "Multiple CT cell types"],
        ],
        header_color="7D3C98")

    doc.add_paragraph()
    add_heading_styled(doc, "9.4 Tissue Response to Injury", 2)

    add_body_text(doc, "The body responds to tissue injury through inflammation and repair:")

    create_styled_table(doc,
        ["Phase", "Time", "Key Events", "Cells Involved"],
        [
            ["Acute Inflammation", "Minutes-hours", "Vasodilation, increased permeability,\nexudate formation, pain, redness, swelling", "Neutrophils (first responders)\nMast cells (histamine)\nMacrophages"],
            ["Chronic Inflammation", "Days-weeks", "Ongoing tissue destruction +\nsimultaneous repair attempts", "Macrophages, Lymphocytes\nPlasma cells\nGiant cells (granulomas)"],
            ["Granulation Tissue", "3-5 days", "New capillaries + fibroblasts +\ninflammatory cells; pink, granular", "Fibroblasts (collagen)\nEndothelial cells (angiogenesis)\nMacrophages"],
            ["Scar Formation", "Weeks-months", "Collagen deposition and maturation\nType III replaced by Type I\nWound contraction (myofibroblasts)", "Fibroblasts\nMyofibroblasts\n(alpha-smooth muscle actin)"],
        ],
        header_color="CB4335")

    doc.add_page_break()


def add_references(doc):
    """Add references section."""
    add_heading_styled(doc, "REFERENCES & RECOMMENDED READING", 1)
    add_section_divider(doc)

    add_numbered_list(doc, [
        "Guyton & Hall: Textbook of Medical Physiology, 14th Edition - Chapters on Tissue Organization",
        "Tortora & Derrickson: Principles of Anatomy & Physiology, 16th Edition - Chapter 4: The Tissue Level of Organization",
        "Ross & Wilson: Anatomy and Physiology in Health and Illness, 14th Edition",
        "Sembulingam K & Sembulingam P: Essentials of Medical Physiology, 8th Edition",
        "Junqueira's Basic Histology: Text and Atlas, 16th Edition (Mescher)",
        "Inderbir Singh: Textbook of Human Histology, 8th Edition",
        "diFiore's Atlas of Histology with Functional Correlations, 13th Edition",
        "Robbins & Cotran: Pathologic Basis of Disease, 10th Edition (for clinical correlations)",
        "Ganong's Review of Medical Physiology, 26th Edition",
        "Netter's Atlas of Human Anatomy (for visual reference)"
    ])

    doc.add_paragraph()
    add_body_text(doc, "Note: These notes are compiled from standard textbooks for exam preparation. "
                  "Always refer to the latest edition of recommended textbooks for your curriculum.")

    doc.add_paragraph()
    doc.add_paragraph()

    # Footer
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("--- END OF NOTES ---")
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x1B, 0x4F, 0x72)

    doc.add_paragraph()
    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = note.add_run("Best of luck for your exams! Study smart, revise often.")
    run.italic = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x6C, 0x3A, 0x83)


def main():
    """Main function to create the complete document."""
    print("Creating 'The Tissues' study notes...")

    # Create document
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # Set margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # Configure heading styles
    for i in range(1, 4):
        heading_style = doc.styles[f'Heading {i}']
        heading_style.font.name = 'Calibri'
        if i == 1:
            heading_style.font.size = Pt(22)
            heading_style.font.bold = True
        elif i == 2:
            heading_style.font.size = Pt(16)
            heading_style.font.bold = True
        elif i == 3:
            heading_style.font.size = Pt(13)
            heading_style.font.bold = True

    # Build document sections
    print("  Adding title page...")
    create_title_page(doc)

    print("  Adding introduction...")
    add_introduction(doc)

    print("  Adding epithelial tissue...")
    add_epithelial_tissue(doc)

    print("  Adding connective tissue...")
    add_connective_tissue(doc)

    print("  Adding muscular tissue...")
    add_muscular_tissue(doc)

    print("  Adding nervous tissue...")
    add_nervous_tissue(doc)

    print("  Adding tissue repair section...")
    add_tissue_repair(doc)

    print("  Adding body membranes...")
    add_membranes_section(doc)

    print("  Adding MCQs...")
    add_mcq_section(doc)

    print("  Adding quick revision...")
    add_quick_revision(doc)

    print("  Adding additional high-yield topics...")
    add_additional_high_yield(doc)

    print("  Adding references...")
    add_references(doc)

    # Save document
    output_path = "The_Tissues_Notes.docx"
    doc.save(output_path)
    print(f"\nDocument saved successfully: {output_path}")
    print("Total sections: 9 (Introduction, Epithelial, Connective, Muscular, Nervous, Repair, Membranes, MCQs, Quick Revision, High-Yield)")
    print("Total MCQs: 50 with detailed explanations")
    print("\nReady for exam preparation!")


if __name__ == "__main__":
    main()
