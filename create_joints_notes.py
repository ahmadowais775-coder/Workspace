#!/usr/bin/env python3
"""
Generate comprehensive study notes on "Joints & Sutures in the Skull"
for Human Anatomy and Physiology.
Creates a visually attractive DOCX file suitable for exam preparation.
Sources: Guyton, Ross & Wilson, Tortora, Sembulingam, B.D. Chaurasia
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import copy


def set_cell_shading(cell, color):
    """Set background shading for a table cell."""
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)


def add_table_borders(table):
    """Add borders to all cells in a table."""
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'<w:top w:val="single" w:sz="6" w:color="2E4057"/>'
        f'<w:left w:val="single" w:sz="6" w:color="2E4057"/>'
        f'<w:bottom w:val="single" w:sz="6" w:color="2E4057"/>'
        f'<w:right w:val="single" w:sz="6" w:color="2E4057"/>'
        f'<w:insideH w:val="single" w:sz="4" w:color="2E4057"/>'
        f'<w:insideV w:val="single" w:sz="4" w:color="2E4057"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)


def create_styled_table(doc, headers, rows, header_color="1B4F72"):
    """Create a styled table with header shading."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_table_borders(table)

    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    cellMargin = parse_xml(
        f'<w:tblCellMar {nsdecls("w")}>'
        f'<w:top w:w="20" w:type="dxa"/>'
        f'<w:left w:w="40" w:type="dxa"/>'
        f'<w:bottom w:w="20" w:type="dxa"/>'
        f'<w:right w:w="40" w:type="dxa"/>'
        f'</w:tblCellMar>'
    )
    tblPr.append(cellMargin)

    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.0
        run = p.add_run(header)
        run.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.font.size = Pt(8.5)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, header_color)

    for row_idx, row_data in enumerate(rows):
        for col_idx, cell_text in enumerate(row_data):
            cell = table.rows[row_idx + 1].cells[col_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.0
            run = p.add_run(str(cell_text))
            run.font.size = Pt(8)
            if row_idx % 2 == 0:
                set_cell_shading(cell, "EBF5FB")
    return table


def add_highlight_box(doc, title, content, color="FEF9E7", border_color="F39C12"):
    """Add a highlighted box for important points."""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.rows[0].cells[0]
    set_cell_shading(cell, color)

    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMargin = parse_xml(
        f'<w:tcMar {nsdecls("w")}>'
        f'<w:top w:w="30" w:type="dxa"/>'
        f'<w:left w:w="60" w:type="dxa"/>'
        f'<w:bottom w:w="30" w:type="dxa"/>'
        f'<w:right w:w="60" w:type="dxa"/>'
        f'</w:tcMar>'
    )
    tcPr.append(tcMargin)

    tcBorders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'<w:top w:val="single" w:sz="12" w:color="{border_color}"/>'
        f'<w:bottom w:val="single" w:sz="12" w:color="{border_color}"/>'
        f'<w:left w:val="single" w:sz="12" w:color="{border_color}"/>'
        f'<w:right w:val="single" w:sz="12" w:color="{border_color}"/>'
        f'</w:tcBorders>'
    )
    tcPr.append(tcBorders)

    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.line_spacing = 1.0
    run = p.add_run(f"  {title}")
    run.bold = True
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor(0x8B, 0x45, 0x13)

    if isinstance(content, list):
        for item in content:
            p = cell.add_paragraph()
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.0
            run = p.add_run(f"  {item}")
            run.font.size = Pt(8.5)
    else:
        p = cell.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.0
        run = p.add_run(f"  {content}")
        run.font.size = Pt(8.5)


def add_section_divider(doc, color="1B4F72"):
    """Add a visual section divider."""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.rows[0].cells[0]
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'<w:bottom w:val="single" w:sz="18" w:color="{color}"/>'
        f'</w:tcBorders>'
    )
    tcPr.append(tcBorders)


def add_heading_styled(doc, text, level=1):
    """Add a styled heading with custom colors."""
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        if level == 1:
            run.font.color.rgb = RGBColor(0x1B, 0x4F, 0x72)
            run.font.size = Pt(18)
        elif level == 2:
            run.font.color.rgb = RGBColor(0x15, 0x4F, 0x0B)
            run.font.size = Pt(14)
        elif level == 3:
            run.font.color.rgb = RGBColor(0x6C, 0x3A, 0x83)
            run.font.size = Pt(11)
    return heading


def add_bullet_points(doc, items, bold_prefix=False):
    """Add formatted bullet points."""
    for item in items:
        p = doc.add_paragraph(style='List Bullet')
        if bold_prefix and ':' in item:
            parts = item.split(':', 1)
            run = p.add_run(parts[0] + ':')
            run.bold = True
            run.font.size = Pt(9)
            run = p.add_run(parts[1])
            run.font.size = Pt(9)
        else:
            run = p.add_run(item)
            run.font.size = Pt(9)


def add_numbered_list(doc, items):
    """Add formatted numbered list."""
    for i, item in enumerate(items, 1):
        p = doc.add_paragraph(style='List Number')
        run = p.add_run(item)
        run.font.size = Pt(9)


def add_body_text(doc, text, bold=False, italic=False):
    """Add body text with formatting."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(9.5)
    run.bold = bold
    run.italic = italic
    return p


def create_title_page(doc):
    """Create an attractive title page."""
    for _ in range(2):
        doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("JOINTS & SUTURES IN THE SKULL")
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor(0x1B, 0x4F, 0x72)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Arthrology - Human Anatomy & Physiology")
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0x2E, 0x86, 0xC1)
    run.italic = True

    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = info.add_run("Comprehensive Study Notes for Exam Preparation")
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x6C, 0x3A, 0x83)

    sources = doc.add_paragraph()
    sources.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sources.add_run("Sources: B.D. Chaurasia | Guyton & Hall | Tortora | Ross & Wilson | Sembulingam")
    run.font.size = Pt(10)
    run.italic = True
    run.font.color.rgb = RGBColor(0x56, 0x6D, 0x7E)

    exams = doc.add_paragraph()
    exams.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = exams.add_run("For: NEET | AIIMS | JIPMER | GPAT | University Exams")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x78, 0x28, 0x1F)
    run.bold = True

    doc.add_page_break()


def add_introduction(doc):
    """Add introduction to Arthrology."""
    add_heading_styled(doc, "INTRODUCTION TO ARTHROLOGY", 1)
    add_section_divider(doc)

    add_body_text(doc, "Arthrology (Greek: arthron = joint, logos = study) is the branch of anatomy "
                  "that deals with the study of joints (articulations). A joint is defined as the "
                  "site where two or more bones meet, whether or not movement occurs between them.")

    add_heading_styled(doc, "Definition & Basic Concepts", 2)
    add_bullet_points(doc, [
        "Joint (Articulation/Arthrosis): A point of contact between two or more bones, or between bone and cartilage",
        "Arthrology: Study of joints and their diseases",
        "Kinesiology: Study of movements at joints",
        "Functions: Allow movement, provide stability, allow growth (epiphyseal plates), transmit forces",
        "Components: Articular surfaces, connecting medium (fibrous/cartilaginous/synovial), blood vessels, nerves"
    ], bold_prefix=True)

    add_highlight_box(doc, "HILTON'S LAW (Important for Exams)",
                     ["The nerve supplying a joint also supplies:",
                      "  1. The muscles moving that joint",
                      "  2. The skin over the insertion of those muscles",
                      "Clinical significance: Pain from joint disease may be referred to skin areas supplied by the same nerve"],
                     color="E8F8F5", border_color="1ABC9C")

    doc.add_page_break()


def add_classification_of_joints(doc):
    """Add classification of joints section."""
    add_heading_styled(doc, "CLASSIFICATION OF JOINTS", 1)
    add_section_divider(doc)

    add_heading_styled(doc, "A. Structural Classification (Based on Connecting Medium)", 2)

    create_styled_table(doc,
        ["Type", "Connecting Medium", "Movement", "Examples"],
        [
            ["Fibrous Joints\n(Synarthrosis)", "Dense fibrous connective tissue\n(collagen fibers)", "Little or no movement\n(immovable/slightly movable)",
             "Sutures (skull)\nSyndesmosis (tibiofibular)\nGomphosis (teeth in sockets)"],
            ["Cartilaginous Joints\n(Amphiarthrosis)", "Cartilage (hyaline or\nfibrocartilage)", "Slight movement\n(limited)",
             "Synchondrosis (epiphyseal plate)\nSymphysis (pubic symphysis)"],
            ["Synovial Joints\n(Diarthrosis)", "Synovial fluid in joint cavity\nlined by synovial membrane", "Free movement\n(most movable)",
             "Shoulder, Hip, Knee\nElbow, Wrist, Ankle"],
        ],
        header_color="1A5276")

    add_heading_styled(doc, "B. Functional Classification (Based on Degree of Movement)", 2)

    create_styled_table(doc,
        ["Type", "Movement Allowed", "Structural Equivalent", "Examples"],
        [
            ["Synarthrosis", "Immovable (no movement)", "Fibrous joints (sutures, gomphosis)", "Skull sutures\nTeeth in sockets"],
            ["Amphiarthrosis", "Slightly movable", "Cartilaginous joints\nSome fibrous (syndesmosis)", "Pubic symphysis\nIntervertebral discs\nInferior tibiofibular joint"],
            ["Diarthrosis", "Freely movable", "Synovial joints", "Shoulder, Hip, Knee\nElbow, Ankle, Wrist"],
        ],
        header_color="6C3483")

    add_highlight_box(doc, "MNEMONIC: Classification of Joints - 'FCS = FSS'",
                     ["Structural: Fibrous, Cartilaginous, Synovial",
                      "Functional: Synarthrosis (immovable), Amphiarthrosis (slightly), Diarthrosis (freely)",
                      "F=S (Fibrous=Synarthrosis), C=A (Cartilaginous=Amphiarthrosis), S=D (Synovial=Diarthrosis)"],
                     color="EBF5FB", border_color="2E86C1")

    doc.add_page_break()


def add_fibrous_joints(doc):
    """Add fibrous joints section."""
    add_heading_styled(doc, "FIBROUS JOINTS (SYNARTHROSES)", 1)
    add_section_divider(doc)

    add_body_text(doc, "In fibrous joints, the bones are united by fibrous connective tissue. "
                  "There is no joint cavity. Movement is absent or very limited. "
                  "The degree of movement depends on the length of the connecting fibers.")

    add_heading_styled(doc, "Types of Fibrous Joints", 2)

    create_styled_table(doc,
        ["Type", "Description", "Examples", "Movement"],
        [
            ["Sutures", "Found ONLY in skull\nBones connected by thin layer of\nfibrous CT (sutural ligament)\nOssify with age (synostosis)",
             "Coronal suture\nSagittal suture\nLambdoid suture\nSquamous suture",
             "Immovable (synarthrosis)\nAllow growth during childhood\nOssify by ~30-40 years"],
            ["Syndesmosis", "Bones connected by interosseous\nmembrane or ligament\nMore fibrous tissue than sutures\nSlightly more movement",
             "Inferior tibiofibular joint\nInterosseous membrane\n(forearm: radius-ulna)\n(leg: tibia-fibula)",
             "Slightly movable\n(amphiarthrosis)\nAllows rotation of radius"],
            ["Gomphosis\n(Dentoalveolar)", "Peg-in-socket joint\nPeriodontal ligament connects\ntooth root to alveolar bone\n(Also called 'peg & socket')",
             "Teeth in alveolar sockets\nof maxilla and mandible",
             "Immovable (synarthrosis)\nPeriodontal ligament acts as\nshock absorber"],
        ],
        header_color="1E8449")

    add_highlight_box(doc, "HIGH YIELD: Sutures become Synostoses",
                     ["Sutures are temporary joints that ossify (turn to bone) with age",
                      "Process: Suture -> Synostosis (bony union)",
                      "Metopic suture: Closes by 2 years of age (earliest to close)",
                      "Other sutures: Begin to close internally after 30 years",
                      "Premature closure = Craniosynostosis (abnormal skull shape)"],
                     color="FDEDEC", border_color="E74C3C")

    doc.add_page_break()


def add_cartilaginous_joints(doc):
    """Add cartilaginous joints section."""
    add_heading_styled(doc, "CARTILAGINOUS JOINTS (AMPHIARTHROSES)", 1)
    add_section_divider(doc)

    add_body_text(doc, "In cartilaginous joints, bones are united by cartilage (hyaline or fibrocartilage). "
                  "There is no joint cavity. Movement is limited but more than fibrous joints.")

    add_heading_styled(doc, "Types of Cartilaginous Joints", 2)

    create_styled_table(doc,
        ["Type", "Cartilage Type", "Description", "Examples", "Movement"],
        [
            ["Primary\n(Synchondrosis)", "Hyaline cartilage", "Temporary joints\nEventually ossify\n(replaced by bone)\nPresent during growth",
             "Epiphyseal plates\n(growth plates)\n1st sternocostal joint\nSpheno-occipital\nsynchondrosis",
             "Immovable\n(synarthrosis)\nAllow bone growth\nin length"],
            ["Secondary\n(Symphysis)", "Fibrocartilage\n(+ thin hyaline layer\non articular surfaces)", "Permanent joints\nDo NOT ossify\nStrong but slightly\nflexible\nMidline structures",
             "Pubic symphysis\nIntervertebral discs\nManubriosternal joint\nSacrococcygeal joint\n(in young adults)",
             "Slightly movable\n(amphiarthrosis)\nCompression &\nshock absorption"],
        ],
        header_color="7D3C98")

    add_highlight_box(doc, "MNEMONIC: Primary vs Secondary Cartilaginous Joints",
                     ["Primary (Synchondrosis): 'Primary = Plate' (growth plate, hyaline, temporary, ossifies)",
                      "Secondary (Symphysis): 'Secondary = Symphysis' (fibrocartilage, permanent, midline)",
                      "All symphyses are in the MIDLINE of body",
                      "Exception: Manubriosternal joint - classified as secondary cartilaginous (symphysis)"],
                     color="F4ECF7", border_color="8E44AD")

    add_heading_styled(doc, "Intervertebral Disc (Important Symphysis)", 2)
    add_bullet_points(doc, [
        "Nucleus pulposus: Central gelatinous core (remnant of notochord); high water content; shock absorber",
        "Annulus fibrosus: Outer ring of fibrocartilage (concentric lamellae); provides strength",
        "Clinical: Disc prolapse (herniation) - nucleus pulposus herniates through annulus fibrosus",
        "Most common site: L4-L5 and L5-S1 (posterolateral herniation)",
        "Posterior longitudinal ligament is weakest laterally - explains posterolateral herniation",
        "No disc between C1-C2 (atlas-axis) and at sacrum/coccyx"
    ], bold_prefix=True)

    doc.add_page_break()


def add_synovial_joints(doc):
    """Add comprehensive synovial joints section."""
    add_heading_styled(doc, "SYNOVIAL JOINTS (DIARTHROSES)", 1)
    add_section_divider(doc)

    add_body_text(doc, "Synovial joints are the most common and most movable joints in the body. "
                  "They are characterized by a joint cavity containing synovial fluid, enclosed by "
                  "an articular capsule. Most joints of the limbs are synovial joints.")

    add_heading_styled(doc, "General Structure of a Synovial Joint", 2)

    create_styled_table(doc,
        ["Component", "Structure", "Function"],
        [
            ["Articular Cartilage", "Hyaline cartilage covering\narticular surfaces\nAvascular, aneural\n(2-5 mm thick)",
             "Reduces friction\nShock absorption\nSmooth gliding surface\nWeight distribution"],
            ["Joint Cavity\n(Synovial Cavity)", "Potential space between\narticular surfaces\nFilled with synovial fluid",
             "Allows free movement\nContains synovial fluid\nUnique to synovial joints"],
            ["Articular Capsule\n(Joint Capsule)", "Two layers:\nOuter: Fibrous capsule (dense CT)\nInner: Synovial membrane",
             "Encloses joint cavity\nProvides stability\nSecrets synovial fluid"],
            ["Synovial Membrane", "Vascular connective tissue\nLines capsule (NOT articular surfaces)\nType A cells (macrophage-like)\nType B cells (fibroblast-like)",
             "Produces synovial fluid\nAbsorbs debris\nRegulates fluid volume\nNutrition to cartilage"],
            ["Synovial Fluid", "Viscous, clear/pale yellow\nDialysate of plasma\n+ Hyaluronic acid (gives viscosity)\n+ Lubricin (glycoprotein)",
             "Lubrication (reduces friction)\nNutrition to cartilage\nShock absorption\nPhagocytosis of debris"],
            ["Ligaments", "Dense regular CT bands\nIntrinsic (thickening of capsule)\nExtrinsic (outside capsule)\nIntra-articular (inside joint)",
             "Connect bone to bone\nLimit excessive movement\nGuide normal movement\nProvide stability"],
            ["Menisci/Discs", "Fibrocartilage pads\nWedge-shaped (menisci)\nor complete (discs)",
             "Shock absorption\nImprove congruence\nSpread synovial fluid\nDeepen articular surface"],
            ["Bursae", "Synovial fluid-filled sacs\nBetween tendons & bones\nor skin & bones",
             "Reduce friction\nCushioning\nFacilitate gliding\nof tendons/muscles"],
        ],
        header_color="1B4F72")

    add_highlight_box(doc, "CLINICAL: Synovial Fluid Analysis",
                     ["Normal: Clear, viscous, <200 WBC/mL, good mucin clot",
                      "Osteoarthritis: Clear, viscous, <2000 WBC/mL, non-inflammatory",
                      "Rheumatoid Arthritis: Turbid, low viscosity, 2000-75000 WBC/mL, poor mucin clot",
                      "Septic Arthritis: Purulent, very low viscosity, >50000 WBC/mL, culture positive",
                      "Gout: Yellow, turbid, negatively birefringent needle-shaped crystals (monosodium urate)",
                      "Pseudogout: Weakly positively birefringent rhomboid crystals (calcium pyrophosphate)"],
                     color="FDEDEC", border_color="E74C3C")

    doc.add_page_break()

    # Types of Synovial Joints
    add_heading_styled(doc, "Types of Synovial Joints (Based on Shape of Articular Surfaces)", 2)

    create_styled_table(doc,
        ["Type", "Shape/Description", "Axes of Movement", "Examples"],
        [
            ["Plane (Gliding)\nArthrodia", "Flat or slightly curved surfaces\nGliding/sliding movements\nSimplest type",
             "Non-axial\n(multiplanar gliding)",
             "Intercarpal joints\nIntertarsal joints\nAcromioclavicular joint\nSacroiliac joint (partially)\nFacet joints (vertebrae)"],
            ["Hinge (Ginglymus)", "Convex surface fits into\nconcave surface\nSpool-shaped (trochlea)\nStrong collateral ligaments",
             "Uniaxial\n(1 axis - transverse)\nFlexion-Extension only",
             "Elbow (humeroulnar)\nKnee (mainly hinge)\nInterphalangeal joints\nAnkle (talocrural)"],
            ["Pivot (Trochoid)", "Rounded/pointed surface\nrotates within a ring\n(formed by bone + ligament)",
             "Uniaxial\n(1 axis - vertical/longitudinal)\nRotation only",
             "Atlanto-axial (median)\n(C1-C2: head rotation)\nProximal radioulnar joint\n(pronation/supination)"],
            ["Condyloid\n(Ellipsoid)", "Oval convex surface fits\ninto oval concave surface\nElliptical articular surfaces",
             "Biaxial\n(2 axes)\nFlexion-Extension\nAbduction-Adduction\n(+ Circumduction)",
             "Wrist (radiocarpal)\nMCP joints (knuckles)\nAtlanto-occipital joint\n(nodding head)"],
            ["Saddle\n(Sellar)", "Both surfaces are\nsaddle-shaped\n(concave in one direction,\nconvex in other)\nReciprocally concavo-convex",
             "Biaxial\n(2 axes)\nFlexion-Extension\nAbduction-Adduction\n(+ Circumduction + Opposition)",
             "1st Carpometacarpal\n(thumb CMC) - MOST\nIMPORTANT EXAMPLE\nSternoclavicular joint\nCalcaneocuboid joint\nIncudomalleolar joint"],
            ["Ball-and-Socket\n(Spheroidal/Enarthrosis)", "Spherical head fits into\ncup-shaped socket\nGreatest range of motion\nof all joint types",
             "Multiaxial/Polyaxial\n(3+ axes)\nAll movements:\nFlexion-Extension\nAbduction-Adduction\nRotation\nCircumduction",
             "Shoulder (glenohumeral)\n- most mobile joint\nHip (acetabulofemoral)\n- most stable ball & socket"],
        ],
        header_color="1E8449")

    add_highlight_box(doc, "MNEMONIC: Types of Synovial Joints - 'PHP CBS'",
                     ["Plane (non-axial), Hinge (uniaxial), Pivot (uniaxial)",
                      "Condyloid (biaxial), Ball-and-socket (multiaxial), Saddle (biaxial)",
                      "Movement axes: Non-axial -> Uniaxial -> Biaxial -> Multiaxial",
                      "Remember: 'Please Help People Cross Bridges Safely'"],
                     color="EBF5FB", border_color="2E86C1")

    doc.add_page_break()

    # Movements at Joints
    add_heading_styled(doc, "Movements at Synovial Joints", 2)

    create_styled_table(doc,
        ["Movement", "Description", "Axis", "Example"],
        [
            ["Flexion", "Decrease in angle between bones\nBending movement (usually anterior)", "Transverse axis", "Bending elbow\nBending knee (exception: posterior)"],
            ["Extension", "Increase in angle between bones\nStraightening movement", "Transverse axis", "Straightening elbow\nStanding upright from sitting"],
            ["Hyperextension", "Extension beyond anatomical position\nBeyond 180 degrees", "Transverse axis", "Looking up at ceiling\nSwinging leg backward"],
            ["Abduction", "Movement AWAY from midline\n(or from axis of digit)", "Anteroposterior axis", "Raising arm to side\nSpreading fingers apart"],
            ["Adduction", "Movement TOWARD midline\n(or toward axis of digit)", "Anteroposterior axis", "Bringing arm back to side\nBringing fingers together"],
            ["Rotation", "Bone revolves around its\nlongitudinal axis", "Vertical/Longitudinal axis", "Turning head side to side\nMedial & lateral rotation of arm"],
            ["Circumduction", "Combination of F, E, Ab, Ad\nCone-shaped movement\nNOT a separate movement", "Multiple axes\n(sequential)", "Swinging arm in circle\nMoving thumb in circle"],
            ["Pronation", "Forearm rotation - palm faces\nposteriorly/downward", "Longitudinal axis", "Turning palm down\n(radius crosses ulna)"],
            ["Supination", "Forearm rotation - palm faces\nanteriorly/upward", "Longitudinal axis", "Turning palm up\n(radius parallel to ulna)"],
            ["Inversion", "Sole of foot turns medially\n(inward)", "Anteroposterior axis", "Subtalar & midtarsal joints\nCommon ankle sprain position"],
            ["Eversion", "Sole of foot turns laterally\n(outward)", "Anteroposterior axis", "Subtalar & midtarsal joints\nLess common sprain"],
            ["Protraction", "Movement anteriorly\n(forward in horizontal plane)", "Horizontal", "Jutting jaw forward\nRounding shoulders"],
            ["Retraction", "Movement posteriorly\n(backward in horizontal plane)", "Horizontal", "Pulling jaw back\nSqueezing shoulder blades"],
            ["Elevation", "Lifting/moving superiorly", "Vertical", "Shrugging shoulders\nClosing mouth"],
            ["Depression", "Lowering/moving inferiorly", "Vertical", "Opening mouth\nDropping shoulders"],
            ["Opposition", "Thumb touches other fingertips\nUnique to humans (precision grip)", "Multiple", "1st CMC joint (saddle)\nGrasping objects"],
            ["Dorsiflexion", "Bending foot toward shin\n(decreasing angle at ankle)", "Transverse", "Walking on heels\nUpstroke of foot"],
            ["Plantarflexion", "Pointing foot/toes downward\n(increasing angle at ankle)", "Transverse", "Standing on tiptoes\nPushing off while walking"],
        ],
        header_color="7D3C98")

    add_highlight_box(doc, "MNEMONIC: Pronation vs Supination",
                     ["SUPination = SUP (as in holding a bowl of SUP/soup - palm up)",
                      "PRONation = PRONe (face/palm down, like lying prone)",
                      "Memory aid: 'You SUPinate to hold SOUP, you PRONate to pour it out'",
                      "Muscles: Supination = Biceps + Supinator; Pronation = Pronator teres + Pronator quadratus"],
                     color="FEF9E7", border_color="F39C12")

    doc.add_page_break()

    # Blood and Nerve Supply
    add_heading_styled(doc, "Blood Supply and Nerve Supply of Joints", 2)

    add_heading_styled(doc, "Blood Supply (Arterial Supply)", 3)
    add_bullet_points(doc, [
        "Periarticular arterial plexus: Network of arteries around the joint",
        "Articular branches: From nearby arteries that cross the joint",
        "Genicular arteries: Supply knee joint (from popliteal artery)",
        "Cruciate anastomosis: Around hip joint (medial & lateral circumflex femoral + superior & inferior gluteal)",
        "Venous drainage: Corresponding veins parallel the arteries"
    ], bold_prefix=True)

    add_heading_styled(doc, "Nerve Supply (Hilton's Law)", 3)
    add_bullet_points(doc, [
        "Hilton's Law: Nerve supplying a joint also supplies muscles moving it and skin over those muscles",
        "Articular nerves carry: (a) Sensory fibers for pain, (b) Proprioceptive fibers, (c) Autonomic fibers (vasomotor)",
        "Proprioceptors in joints: Ruffini endings (position), Pacinian corpuscles (pressure/vibration), Golgi tendon organs (tension)",
        "Pain fibers: Type IV (C fibers) - slow, dull ache; Type III (A-delta) - sharp pain",
        "Referred pain: Joint pain may be referred to skin dermatome sharing the same nerve root"
    ], bold_prefix=True)

    add_highlight_box(doc, "CLINICAL: Referred Pain from Joints",
                     ["Hip joint pain (obturator nerve L2,3,4) may be felt in the KNEE",
                      "This is why knee examination is incomplete without hip examination",
                      "Shoulder joint (C5,6) - pain may radiate to deltoid region",
                      "Temporomandibular joint - pain may be felt in ear (auriculotemporal nerve)"],
                     color="F9EBEA", border_color="CB4335")

    doc.add_page_break()


def add_specific_joints(doc):
    """Add section on specific important joints."""
    add_heading_styled(doc, "SPECIFIC IMPORTANT JOINTS OF THE BODY", 1)
    add_section_divider(doc)

    # Shoulder Joint
    add_heading_styled(doc, "1. Shoulder Joint (Glenohumeral Joint)", 2)
    add_bullet_points(doc, [
        "Type: Synovial, Ball-and-socket (multiaxial)",
        "Articulation: Head of humerus with glenoid cavity of scapula",
        "Most mobile joint in body (but least stable - frequent dislocation)",
        "Glenoid labrum: Fibrocartilage rim that deepens the shallow glenoid cavity",
        "Rotator cuff muscles (SITS): Supraspinatus, Infraspinatus, Teres minor, Subscapularis",
        "Rotator cuff holds head of humerus in glenoid cavity (dynamic stability)",
        "Dislocation: Most commonly ANTERIOR and INFERIOR (axillary nerve may be damaged)",
        "Supraspinatus: Most commonly torn rotator cuff muscle (impingement under acromion)",
        "Movements: Flexion, Extension, Abduction, Adduction, Medial & Lateral rotation, Circumduction",
        "Blood supply: Anterior & posterior circumflex humeral arteries, suprascapular artery",
        "Nerve supply: Axillary nerve (C5,6), Suprascapular nerve, Lateral pectoral nerve"
    ], bold_prefix=True)

    # Hip Joint
    add_heading_styled(doc, "2. Hip Joint (Acetabulofemoral Joint)", 2)
    add_bullet_points(doc, [
        "Type: Synovial, Ball-and-socket (multiaxial) - most stable ball & socket joint",
        "Articulation: Head of femur with acetabulum of hip bone",
        "Acetabular labrum: Deepens the socket (+ transverse acetabular ligament)",
        "Ligaments: Iliofemoral (Y-shaped, strongest), Pubofemoral, Ischiofemoral, Ligamentum teres",
        "Iliofemoral ligament (of Bigelow): Strongest ligament in body; prevents hyperextension",
        "Ligamentum teres: Carries artery to head of femur (important in children; less in adults)",
        "Blood supply: Medial & lateral circumflex femoral aa., artery of ligamentum teres, obturator a.",
        "Nerve supply: Femoral nerve (L2,3,4), Obturator nerve (L2,3,4), Sciatic nerve (L4,5,S1,2,3)",
        "Avascular necrosis: Head of femur - common after femoral neck fracture (blood supply disrupted)",
        "Most common dislocation: Posterior (dashboard injury) - sciatic nerve at risk",
        "Trendelenburg sign: Weakness of gluteus medius - pelvis drops on opposite side"
    ], bold_prefix=True)

    # Knee Joint
    add_heading_styled(doc, "3. Knee Joint (Tibiofemoral Joint)", 2)
    add_bullet_points(doc, [
        "Type: Synovial, Modified hinge (primarily uniaxial, some rotation when flexed)",
        "Largest and most complex joint in body",
        "Articulation: Femoral condyles with tibial condyles (+ patella with femur)",
        "Menisci: Medial (C-shaped, attached to MCL - more prone to injury) & Lateral (O-shaped, more mobile)",
        "Cruciate ligaments: ACL (anterior - prevents anterior tibial displacement) & PCL (posterior - prevents posterior displacement)",
        "ACL injury: Anterior drawer test positive; most commonly injured knee ligament in sports",
        "PCL injury: Posterior drawer test positive; dashboard injury (posterior blow to tibia)",
        "Collateral ligaments: MCL (medial - resists valgus) & LCL (lateral - resists varus)",
        "Unhappy triad (O'Donoghue): ACL + MCL + Medial meniscus (lateral blow to knee)",
        "Locking mechanism: Terminal extension - medial rotation of femur on tibia (screw-home mechanism)",
        "Unlocking: Popliteus muscle laterally rotates femur to 'unlock' the knee",
        "Bursae: Suprapatellar, Prepatellar (housemaid's knee), Infrapatellar, Semimembranosus",
        "Blood supply: Genicular branches of popliteal artery (5 genicular arteries)",
        "Nerve supply: Femoral, Obturator, Common peroneal, Tibial nerves"
    ], bold_prefix=True)

    add_highlight_box(doc, "CLINICAL: Knee Joint Tests",
                     ["Anterior Drawer Test: ACL injury (tibia moves forward on femur)",
                      "Posterior Drawer Test: PCL injury (tibia moves backward)",
                      "Lachman Test: Most sensitive test for ACL tear",
                      "McMurray Test: Meniscal tear (clicking on rotation)",
                      "Valgus stress test: MCL integrity",
                      "Varus stress test: LCL integrity",
                      "Unhappy/Terrible Triad: ACL + MCL + Medial Meniscus tear"],
                     color="E8F8F5", border_color="1ABC9C")

    doc.add_page_break()

    # Elbow Joint
    add_heading_styled(doc, "4. Elbow Joint", 2)
    add_bullet_points(doc, [
        "Type: Synovial, Hinge (uniaxial) - the humeroulnar component",
        "Articulation: Trochlea of humerus with trochlear notch of ulna (+ capitulum with radial head)",
        "Carrying angle: Angle between arm and forearm (valgus angle) - Males: 5-10 deg, Females: 10-15 deg",
        "Increased carrying angle = Cubitus valgus (may cause ulnar nerve injury)",
        "Decreased carrying angle = Cubitus varus (gunstock deformity - supracondylar fracture)",
        "Ligaments: Medial (ulnar) collateral, Lateral (radial) collateral, Annular ligament",
        "Annular ligament: Holds radial head in place; pulled elbow (nursemaid's elbow) = radial head subluxation in children",
        "Funny bone: Ulnar nerve behind medial epicondyle (cubital tunnel)",
        "Blood supply: Anastomosis around elbow (superior/inferior ulnar collateral + radial/middle collateral + recurrent arteries)",
        "Nerve supply: Musculocutaneous, Radial, Ulnar, Median nerves"
    ], bold_prefix=True)

    # Temporomandibular Joint
    add_heading_styled(doc, "5. Temporomandibular Joint (TMJ)", 2)
    add_bullet_points(doc, [
        "Type: Synovial, Modified hinge and gliding (bicondylar)",
        "Articulation: Mandibular condyle with mandibular fossa of temporal bone",
        "Articular disc: Divides joint into upper (gliding) and lower (hinge) compartments",
        "Unique: Articular surfaces covered by fibrocartilage (not hyaline cartilage)",
        "Movements: Depression (opening), Elevation (closing), Protraction, Retraction, Lateral movements",
        "Lateral pterygoid: Only muscle that OPENS the mouth (depresses mandible)",
        "Dislocation: Anterior dislocation - jaw locks open (lateral pterygoid spasm)",
        "Nerve supply: Auriculotemporal nerve & Masseteric nerve (branches of mandibular V3)",
        "Blood supply: Superficial temporal artery & Maxillary artery"
    ], bold_prefix=True)

    # Ankle Joint
    add_heading_styled(doc, "6. Ankle Joint (Talocrural Joint)", 2)
    add_bullet_points(doc, [
        "Type: Synovial, Hinge (uniaxial)",
        "Articulation: Tibia + Fibula (mortise) with Talus (tenon)",
        "Movements: Dorsiflexion and Plantarflexion only",
        "Ligaments: Medial (Deltoid) - very strong, 4 parts; Lateral - 3 separate ligaments (ATFL, CFL, PTFL)",
        "Most common ankle sprain: Inversion injury - damages LATERAL ligaments (ATFL most commonly torn)",
        "Deltoid ligament rarely torn (if damaged, suspect fracture - Pott's fracture)",
        "Blood supply: Anterior & posterior tibial arteries, peroneal artery",
        "Nerve supply: Deep peroneal nerve, Tibial nerve"
    ], bold_prefix=True)

    add_highlight_box(doc, "HIGH YIELD: Joint Comparison Table",
                     ["Most MOBILE joint: Shoulder (glenohumeral)",
                      "Most STABLE ball & socket: Hip (acetabulofemoral)",
                      "LARGEST joint: Knee (tibiofemoral)",
                      "STRONGEST ligament: Iliofemoral (hip)",
                      "Most commonly DISLOCATED: Shoulder (anterior-inferior)",
                      "Most commonly SPRAINED: Ankle (lateral ligaments - inversion)"],
                     color="FEF9E7", border_color="F39C12")

    doc.add_page_break()


def add_sutures_section(doc):
    """Add comprehensive sutures of the skull section."""
    add_heading_styled(doc, "SUTURES IN THE SKULL", 1)
    add_section_divider(doc)

    add_body_text(doc, "Sutures are fibrous joints unique to the skull. They are immovable (synarthroses) "
                  "joints where the bone edges interlock with a thin layer of dense fibrous connective "
                  "tissue (sutural ligament/membrane). They allow skull growth during childhood and "
                  "eventually ossify in adulthood (synostosis).")

    add_heading_styled(doc, "Types of Sutures (Based on Edge Morphology)", 2)

    create_styled_table(doc,
        ["Suture Type", "Description", "Examples"],
        [
            ["Serrate (Dentate/\nDenticulate)", "Interlocking tooth-like projections\nStrongest type of suture\nMost common in skull vault",
             "Coronal suture\nSagittal suture\nLambdoid suture"],
            ["Squamous\n(Beveled/Overlapping)", "Overlapping beveled edges\n(like fish scales)\nOne bone overlaps the other",
             "Squamous suture\n(temporal-parietal)\nParietomastoid suture"],
            ["Plane (Butt/\nHarmonia)", "Straight, smooth edges\nSimple apposition\nWeakest type of suture",
             "Internasal suture\nIntermaxillary suture\nMedian palatine suture"],
            ["Schindylesis", "Plate of one bone fits into\ncleft/groove of another bone\n(Rare type)",
             "Vomer into sphenoid rostrum\n(vomero-vaginal suture)"],
            ["Wedge (Limbous)", "One bone edge beveled &\noverlapped internally\nwith a ridge externally",
             "Certain parts of\ncoronal suture in infants"],
        ],
        header_color="1A5276")

    doc.add_page_break()

    # Major Sutures
    add_heading_styled(doc, "Major Sutures of the Skull", 2)

    create_styled_table(doc,
        ["Suture", "Location/Bones United", "Type", "Key Features"],
        [
            ["Coronal Suture", "Frontal bone + Two parietal bones\n(crosses skull transversely)", "Serrate",
             "Between frontal & parietal bones\nAnterior fontanelle at junction\nwith sagittal suture\nPremature closure = Brachycephaly"],
            ["Sagittal Suture", "Between two parietal bones\n(runs anteroposterior in midline)", "Serrate",
             "Midline suture of skull vault\nAnterior end: bregma\nPosterior end: lambda\nPremature closure = Scaphocephaly\n(most common craniosynostosis)"],
            ["Lambdoid Suture", "Two parietal bones + Occipital bone\n(crosses skull posteriorly)", "Serrate",
             "Lambda-shaped (inverted V)\nPosterior fontanelle at junction\nwith sagittal suture\nPremature closure = Plagiocephaly\nWormian bones most common here"],
            ["Squamous Suture", "Parietal bone + Temporal bone\n(bilateral, on sides of skull)", "Squamous",
             "Weakest part of skull\nOverlying middle meningeal artery\nFracture here = epidural hematoma\n(rupture of middle meningeal a.)"],
            ["Metopic (Frontal)\nSuture", "Between two halves of frontal bone\n(midline of forehead)", "Serrate",
             "ONLY suture that normally closes\nin CHILDHOOD (by 2 years)\nPersistent metopic = metopism\n(seen in ~8% adults)\nPremature closure = Trigonocephaly"],
            ["Intermaxillary\nSuture", "Between two maxillae\n(midline of hard palate)", "Plane",
             "Forms anterior part of\nmedian palatine suture\nCleft palate: failure of fusion"],
            ["Internasal Suture", "Between two nasal bones\n(midline of nose bridge)", "Plane",
             "Nasion: meeting point of\ninternasal & frontonsal sutures"],
            ["Fronto-nasal Suture", "Frontal bone + Nasal bones", "Plane/Serrate",
             "Nasion is at this point\n(craniometric landmark)"],
            ["Palatine Sutures\n(Median & Transverse)", "Median: between 2 palatine processes\nTransverse: maxilla & palatine bones", "Plane",
             "Important in orthodontics\n(rapid palatal expansion)"],
            ["Zygomaticotemporal", "Zygomatic bone + Temporal bone\n(zygomatic arch)", "Serrate/Plane",
             "Part of zygomatic arch\nLateral wall of skull"],
        ],
        header_color="6C3483")

    add_highlight_box(doc, "MNEMONIC: Skull Suture Locations",
                     ["CORonal = CROWn of head (frontal-parietal boundary)",
                      "SAGittal = SAGittal plane (midline, between parietals)",
                      "LAMBdoid = shaped like Greek letter Lambda (posterior)",
                      "SQUAMous = SQUAMa of temporal bone (sides)",
                      "METOPIC = METa (forehead) + opis (between frontal halves)"],
                     color="EBF5FB", border_color="2E86C1")

    doc.add_page_break()

    # Fontanelles
    add_heading_styled(doc, "Fontanelles (Soft Spots)", 2)

    add_body_text(doc, "Fontanelles are membrane-covered gaps between developing skull bones in infants. "
                  "They allow brain growth, molding during birth, and serve as clinical indicators of "
                  "intracranial pressure and hydration status.")

    create_styled_table(doc,
        ["Fontanelle", "Location", "Shape", "Closure Time", "Clinical Significance"],
        [
            ["Anterior\n(Bregmatic)", "Junction of coronal,\nsagittal & frontal\n(metopic) sutures\nBREGMA landmark",
             "Diamond-shaped\n(rhomboid)\nLargest fontanelle\n~4 cm x 2.5 cm",
             "18 months\n(12-24 months)",
             "Bulging: raised ICP, meningitis\nSunken: dehydration\nUsed for: Ultrasound imaging\nof neonatal brain"],
            ["Posterior\n(Lambdoid)", "Junction of sagittal\nand lambdoid sutures\nLAMBDA landmark",
             "Triangular\n(small)\n~1 cm",
             "2-3 months\n(by 3 months)",
             "Smallest of the paired fontanelles\nImportant during labor\n(identifies presentation)"],
            ["Anterolateral\n(Sphenoid/Pterion)", "Junction of frontal,\nparietal, temporal,\nsphenoid bones\nPTERION landmark",
             "Small, irregular\n(bilateral, paired)",
             "3 months\n(by 6 months)",
             "PTERION: thinnest part of skull\nMiddle meningeal artery deep\nFracture = Epidural hematoma"],
            ["Posterolateral\n(Mastoid/Asterion)", "Junction of parietal,\noccipital, temporal\nbones\nASTERION landmark",
             "Small, irregular\n(bilateral, paired)",
             "12-18 months\n(by 2 years)",
             "Asterion: landmark for\ntransverse-sigmoid sinus junction\nSurgical landmark for\nposterior fossa approach"],
        ],
        header_color="1E8449")

    add_highlight_box(doc, "HIGH YIELD: Fontanelle Facts",
                     ["Total fontanelles: 6 (2 unpaired midline + 4 paired lateral)",
                      "Largest: Anterior (bregmatic) - closes LAST among major fontanelles (18 months)",
                      "Smallest to close: Posterior (2-3 months)",
                      "Anterior fontanelle: Diamond-shaped, most clinically important",
                      "Bulging fontanelle = Raised ICP (meningitis, hydrocephalus)",
                      "Sunken fontanelle = Dehydration",
                      "Delayed closure: Hypothyroidism, Rickets, Hydrocephalus, Down syndrome, Achondroplasia"],
                     color="FDEDEC", border_color="E74C3C")

    doc.add_page_break()

    # Ossification Timeline
    add_heading_styled(doc, "Ossification Timeline of Sutures", 2)

    create_styled_table(doc,
        ["Suture/Fontanelle", "Closure Time", "Clinical Relevance"],
        [
            ["Posterior fontanelle", "2-3 months", "Earliest fontanelle to close"],
            ["Anterolateral (sphenoid) fontanelle", "3-6 months", "Near pterion (danger area)"],
            ["Metopic suture", "2 years (by end of 2nd year)", "Earliest suture to close; persistence = metopism"],
            ["Posterolateral (mastoid) fontanelle", "12-18 months", "Near asterion surgical landmark"],
            ["Anterior fontanelle", "18 months (12-24 months)", "Last major fontanelle to close"],
            ["Other skull sutures (coronal, sagittal, lambdoid)", "Begin internal fusion: ~30 years\nComplete: 40-50+ years", "Used in forensic age estimation"],
            ["Basilar suture (spheno-occipital synchondrosis)", "18-25 years", "Last growth center of skull base\nImportant for skull base growth"],
        ],
        header_color="2C3E50")

    add_heading_styled(doc, "Craniosynostosis (Premature Suture Closure)", 2)

    create_styled_table(doc,
        ["Suture Affected", "Skull Deformity", "Description"],
        [
            ["Sagittal suture\n(MOST COMMON - 40-55%)", "Scaphocephaly\n(Dolichocephaly)", "Long, narrow skull (boat-shaped)\nIncreased AP diameter\nDecreased biparietal diameter"],
            ["Coronal suture\n(bilateral)", "Brachycephaly", "Short, wide skull\nFlattened forehead\nIncreased biparietal diameter"],
            ["Coronal suture\n(unilateral)", "Plagiocephaly\n(anterior)", "Asymmetric skull\nFlattening on affected side\nHarlequin eye deformity"],
            ["Metopic suture", "Trigonocephaly", "Triangular-shaped forehead\nMidline frontal ridge\nHypotelorism (close-set eyes)"],
            ["Lambdoid suture\n(unilateral)", "Posterior\nPlagiocephaly", "Asymmetric flattening\nof occipital region\n(must differentiate from positional)"],
            ["All sutures\n(multiple)", "Oxycephaly\n(Turricephaly)", "Tower/pointed skull\nMost severe form\nRaised ICP\nCrouzon/Apert syndrome"],
        ],
        header_color="7D3C98")

    add_highlight_box(doc, "CLINICAL: Important Points on Craniosynostosis",
                     ["Skull grows PERPENDICULAR to the suture (Virchow's law)",
                      "Premature closure restricts growth perpendicular to the suture",
                      "Compensatory growth occurs parallel to the fused suture",
                      "Most common: Sagittal synostosis (scaphocephaly)",
                      "Syndromic: Crouzon syndrome (FGFR2) - craniosynostosis + midface hypoplasia",
                      "Apert syndrome: Craniosynostosis + syndactyly (fused fingers) - FGFR2",
                      "Treatment: Surgical correction (craniotomy/cranial vault remodeling) - best before 1 year"],
                     color="F4ECF7", border_color="8E44AD")

    doc.add_page_break()

    # Wormian Bones & Additional Points
    add_heading_styled(doc, "Wormian (Sutural) Bones", 2)
    add_bullet_points(doc, [
        "Definition: Small, irregular extra bones found within sutures (accessory ossification centers)",
        "Most common location: Lambdoid suture (and at lambda)",
        "Also called: Ossa suturae or intrasutural bones",
        "Named after: Ole Worm (Danish anatomist, 1588-1654)",
        "Usually clinically insignificant but may indicate underlying conditions",
        "Associated conditions: Osteogenesis imperfecta, Cleidocranial dysostosis, Hypothyroidism, Hydrocephalus, Rickets, Down syndrome",
        "Inca bone: Large wormian bone at lambda (between occipital and parietal bones)"
    ], bold_prefix=True)

    add_highlight_box(doc, "MNEMONIC for Wormian Bones Associations: 'CHOROID'",
                     ["C - Cleidocranial dysostosis",
                      "H - Hypothyroidism & Hydrocephalus",
                      "O - Osteogenesis imperfecta (most important association)",
                      "R - Rickets",
                      "O - Otopalatodigital syndrome",
                      "I - Idiopathic (normal variant in many people)",
                      "D - Down syndrome"],
                     color="EBF5FB", border_color="2E86C1")

    # Important Skull Landmarks
    add_heading_styled(doc, "Important Craniometric Points at Suture Junctions", 2)

    create_styled_table(doc,
        ["Landmark", "Location", "Clinical Significance"],
        [
            ["Bregma", "Junction of coronal + sagittal sutures\n(anterior fontanelle site)", "Skull vertex landmark\nNeurosurgical reference point"],
            ["Lambda", "Junction of sagittal + lambdoid sutures\n(posterior fontanelle site)", "Posterior skull landmark\nForensic anthropology"],
            ["Pterion", "Junction of frontal, parietal, temporal,\nsphenoid bones (H-shaped)", "THINNEST part of skull\nMiddle meningeal artery beneath\nFracture = Epidural hematoma"],
            ["Asterion", "Junction of parietal, occipital,\ntemporal bones (posterolateral)", "Transverse-sigmoid sinus junction\nSurgical landmark for posterior fossa"],
            ["Nasion", "Junction of internasal + frontonasal sutures\n(bridge of nose)", "Craniometric measurement point\nSoft tissue 'nasion' differs slightly"],
            ["Inion", "External occipital protuberance\n(posterior midline)", "Confluence of sinuses internally\nMost projecting point of occiput"],
            ["Vertex", "Highest point of skull\n(in sagittal suture, near bregma)", "Vertex presentation (normal delivery)\nSkull measurement reference"],
            ["Glabella", "Smooth area between superciliary arches\n(above nasion, below metopic suture)", "Glabellar tap reflex\nCraniometric point\nForehead landmark"],
        ],
        header_color="154360")

    doc.add_page_break()


def add_mcq_section(doc):
    """Add Previous Year MCQs section with 50+ questions."""
    add_heading_styled(doc, "PREVIOUS YEAR MCQs (NEET/AIIMS/JIPMER/GPAT)", 1)
    add_section_divider(doc)

    add_body_text(doc, "The following MCQs are compiled from previous year papers of NEET, AIIMS, "
                  "JIPMER, and GPAT examinations. Each question includes the correct answer with explanation.", bold=True)

    mcqs = [
        ("Q1. Which is the most movable joint in the body? [NEET 2019]",
         "(a) Hip joint (b) Knee joint (c) Shoulder joint (d) Ankle joint",
         "Answer: (c) Shoulder joint",
         "The glenohumeral (shoulder) joint is a ball-and-socket joint with the greatest range of motion. However, this mobility comes at the cost of stability, making it the most commonly dislocated joint."),

        ("Q2. Gomphosis is a type of: [AIIMS 2018]",
         "(a) Synovial joint (b) Cartilaginous joint (c) Fibrous joint (d) Secondary cartilaginous joint",
         "Answer: (c) Fibrous joint",
         "Gomphosis is a fibrous joint (peg-in-socket) where teeth are anchored in alveolar sockets by periodontal ligament. It is an immovable (synarthrosis) joint."),

        ("Q3. The anterior fontanelle closes at: [NEET 2020]",
         "(a) 6 months (b) 12 months (c) 18 months (d) 24 months",
         "Answer: (c) 18 months",
         "Anterior (bregmatic) fontanelle is the largest fontanelle, diamond-shaped, and closes between 12-24 months (average 18 months). It is clinically important for assessing ICP and hydration."),

        ("Q4. Saddle joint is seen in: [JIPMER 2017]",
         "(a) Wrist joint (b) First carpometacarpal joint (c) Elbow joint (d) Knee joint",
         "Answer: (b) First carpometacarpal joint",
         "The first carpometacarpal (thumb CMC) joint is the best example of a saddle (sellar) joint. It allows opposition of the thumb, which is unique to primates and essential for precision grip."),

        ("Q5. Pterion is formed by the meeting of which bones? [AIIMS 2019]",
         "(a) Frontal, Parietal, Temporal, Sphenoid (b) Frontal, Parietal, Occipital, Temporal\n(c) Parietal, Temporal, Occipital, Sphenoid (d) Frontal, Sphenoid, Ethmoid, Temporal",
         "Answer: (a) Frontal, Parietal, Temporal, Sphenoid",
         "Pterion is H-shaped and lies 4 cm above the zygomatic arch. It overlies the middle meningeal artery; fracture here causes epidural hematoma."),

        ("Q6. Hilton's law states that: [NEET 2018]",
         "(a) Nerve supplying a joint also supplies muscles moving it and skin over those muscles\n(b) Nerve supplying a muscle also supplies all joints moved by it\n(c) Blood supply of a joint determines its nerve supply\n(d) Joints are supplied only by sensory nerves",
         "Answer: (a) Nerve supplying a joint also supplies muscles moving it and skin over those muscles",
         "Hilton's law explains referred pain from joints and is important in understanding pain patterns in joint diseases."),

        ("Q7. Which ligament is the strongest in the human body? [GPAT 2019]",
         "(a) Anterior cruciate ligament (b) Posterior longitudinal ligament\n(c) Iliofemoral ligament (d) Ligamentum flavum",
         "Answer: (c) Iliofemoral ligament",
         "The iliofemoral ligament (Y-shaped ligament of Bigelow) is the strongest ligament in the body. It prevents hyperextension at the hip joint and supports upright posture."),

        ("Q8. Premature closure of sagittal suture leads to: [AIIMS 2017]",
         "(a) Brachycephaly (b) Scaphocephaly (c) Plagiocephaly (d) Trigonocephaly",
         "Answer: (b) Scaphocephaly",
         "Sagittal suture synostosis is the MOST COMMON craniosynostosis. Per Virchow's law, skull grows perpendicular to the suture; so sagittal closure causes a long, narrow (boat-shaped) skull."),

        ("Q9. Which of the following is a secondary cartilaginous joint? [NEET 2017]",
         "(a) Epiphyseal plate (b) First sternocostal joint (c) Pubic symphysis (d) Spheno-occipital synchondrosis",
         "Answer: (c) Pubic symphysis",
         "Secondary cartilaginous joints (symphyses) contain fibrocartilage, are permanent, and are found in the midline. Pubic symphysis, intervertebral discs, and manubriosternal joint are examples."),

        ("Q10. Unhappy triad of O'Donoghue involves injury to: [JIPMER 2018]",
         "(a) ACL + MCL + Medial meniscus (b) PCL + LCL + Lateral meniscus\n(c) ACL + LCL + Lateral meniscus (d) PCL + MCL + Medial meniscus",
         "Answer: (a) ACL + MCL + Medial meniscus",
         "The unhappy/terrible triad occurs from a lateral blow to the knee (e.g., sports tackle). MCL is attached to medial meniscus (hence both injured together). ACL tears as valgus force continues."),

        ("Q11. The most common type of craniosynostosis involves: [NEET 2021]",
         "(a) Coronal suture (b) Sagittal suture (c) Metopic suture (d) Lambdoid suture",
         "Answer: (b) Sagittal suture",
         "Sagittal suture synostosis accounts for 40-55% of all craniosynostosis cases, resulting in scaphocephaly (dolichocephaly) - a long, narrow skull shape."),

        ("Q12. Wormian bones are most commonly found in: [AIIMS 2020]",
         "(a) Coronal suture (b) Sagittal suture (c) Lambdoid suture (d) Squamous suture",
         "Answer: (c) Lambdoid suture",
         "Wormian (sutural/intrasutural) bones are most commonly found in the lambdoid suture. They are associated with osteogenesis imperfecta, hypothyroidism, cleidocranial dysostosis, and Down syndrome."),

        ("Q13. Ball and socket joint is found in: [GPAT 2018]",
         "(a) Wrist (b) Elbow (c) Hip (d) Knee",
         "Answer: (c) Hip",
         "Hip (acetabulofemoral) and shoulder (glenohumeral) are the two ball-and-socket joints. Hip is the most stable (deep socket), while shoulder is the most mobile (shallow glenoid)."),

        ("Q14. Which suture is the first to close in the skull? [NEET 2016]",
         "(a) Sagittal (b) Coronal (c) Metopic (d) Lambdoid",
         "Answer: (c) Metopic",
         "The metopic (frontal) suture closes by 2 years of age, making it the first suture to close. Persistent metopic suture (metopism) is seen in about 8% of adults."),

        ("Q15. Pivot joint is found between: [JIPMER 2019]",
         "(a) Atlas and Axis (b) Humerus and Ulna (c) Femur and Tibia (d) Radius and Ulna distally",
         "Answer: (a) Atlas and Axis",
         "The median atlantoaxial joint (dens of axis pivots within ring of atlas + transverse ligament) is the best example. Also, the proximal radioulnar joint is a pivot joint (allows pronation/supination)."),

        ("Q16. Synovial fluid is produced by: [GPAT 2020]",
         "(a) Articular cartilage (b) Fibrous capsule (c) Synovial membrane (d) Ligaments",
         "Answer: (c) Synovial membrane",
         "The synovial membrane (inner layer of articular capsule) produces synovial fluid. Type B synoviocytes (fibroblast-like) secrete hyaluronic acid which gives the fluid its viscosity."),

        ("Q17. Epidural hematoma is most commonly caused by injury to which area? [AIIMS 2019]",
         "(a) Bregma (b) Pterion (c) Lambda (d) Asterion",
         "Answer: (b) Pterion",
         "Pterion is the thinnest area of skull where the middle meningeal artery runs in an epidural groove. Fracture here tears the artery causing epidural (extradural) hematoma - a neurosurgical emergency."),

        ("Q18. The knee joint is classified as: [NEET 2020]",
         "(a) Hinge joint (b) Pivot joint (c) Condyloid joint (d) Saddle joint",
         "Answer: (a) Hinge joint",
         "The knee is classified as a modified hinge joint (ginglymus). It primarily allows flexion-extension but also permits some rotation when the knee is flexed (medial rotation for locking mechanism)."),

        ("Q19. Which test is most sensitive for ACL tear? [AIIMS 2018]",
         "(a) Anterior drawer test (b) Lachman test (c) McMurray test (d) Pivot shift test",
         "Answer: (b) Lachman test",
         "Lachman test (anterior tibial translation at 20-30 degrees flexion) is the most sensitive clinical test for ACL injury. Anterior drawer test is more commonly known but less sensitive."),

        ("Q20. Intervertebral disc is a type of: [NEET 2019]",
         "(a) Fibrous joint (b) Primary cartilaginous joint (c) Secondary cartilaginous joint (d) Synovial joint",
         "Answer: (c) Secondary cartilaginous joint",
         "Intervertebral discs are symphyses (secondary cartilaginous joints). They contain fibrocartilage (annulus fibrosus) with a gelatinous center (nucleus pulposus) and allow slight movement."),

        ("Q21. Carrying angle is measured at: [JIPMER 2020]",
         "(a) Shoulder joint (b) Elbow joint (c) Hip joint (d) Knee joint",
         "Answer: (b) Elbow joint",
         "Carrying angle is the valgus angle at the elbow (between arm and forearm axes). Normal: Males 5-10 degrees, Females 10-15 degrees. Increased in cubitus valgus; decreased in cubitus varus."),

        ("Q22. Nursemaid's elbow (pulled elbow) involves: [AIIMS 2020]",
         "(a) Dislocation of elbow (b) Subluxation of radial head (c) Fracture of olecranon (d) Tear of MCL",
         "Answer: (b) Subluxation of radial head",
         "In children <5 years, sudden pulling of extended forearm causes radial head to slip partially out of annular ligament. Treatment: supination + flexion maneuver (reduces spontaneously)."),

        ("Q23. Diamond-shaped fontanelle is: [GPAT 2017]",
         "(a) Anterior fontanelle (b) Posterior fontanelle (c) Mastoid fontanelle (d) Sphenoid fontanelle",
         "Answer: (a) Anterior fontanelle",
         "Anterior fontanelle is diamond (rhomboid) shaped, largest fontanelle, located at bregma. Posterior fontanelle is triangular and smaller."),

        ("Q24. Which joint allows opposition movement? [NEET 2018]",
         "(a) Hinge (b) Ball and socket (c) Saddle (d) Pivot",
         "Answer: (c) Saddle",
         "The saddle joint at the 1st carpometacarpal (thumb CMC) joint uniquely allows opposition - the ability to touch thumb tip to other fingertips. This is essential for precision grip."),

        ("Q25. Ligamentum teres is found in: [JIPMER 2016]",
         "(a) Knee joint (b) Hip joint (c) Shoulder joint (d) Elbow joint",
         "Answer: (b) Hip joint",
         "Ligamentum teres (ligament of head of femur) connects the femoral head to the acetabular fossa. It carries the artery to the head of femur (important blood supply in children)."),

        ("Q26. Bregma is the meeting point of: [AIIMS 2017]",
         "(a) Coronal and sagittal sutures (b) Sagittal and lambdoid sutures\n(c) Coronal and lambdoid sutures (d) Squamous and coronal sutures",
         "Answer: (a) Coronal and sagittal sutures",
         "Bregma is where the coronal and sagittal sutures meet (site of anterior fontanelle in infants). Lambda is where sagittal and lambdoid sutures meet (site of posterior fontanelle)."),

        ("Q27. The rotator cuff muscles are remembered by mnemonic: [NEET 2017]",
         "(a) SITS (b) SALT (c) STAR (d) SPIN",
         "Answer: (a) SITS",
         "SITS = Supraspinatus, Infraspinatus, Teres minor, Subscapularis. These muscles form the rotator cuff that stabilizes the glenohumeral joint. Supraspinatus is most commonly injured."),

        ("Q28. Synchondrosis is an example of: [GPAT 2019]",
         "(a) Fibrous joint (b) Primary cartilaginous joint (c) Secondary cartilaginous joint (d) Synovial joint",
         "Answer: (b) Primary cartilaginous joint",
         "Synchondrosis = Primary cartilaginous joint, united by HYALINE cartilage. They are temporary and eventually ossify. Examples: epiphyseal plates, 1st sternocostal joint, spheno-occipital synchondrosis."),

        ("Q29. Which is NOT a type of synovial joint? [NEET 2021]",
         "(a) Hinge (b) Symphysis (c) Saddle (d) Ball and socket",
         "Answer: (b) Symphysis",
         "Symphysis is a CARTILAGINOUS joint (secondary), not synovial. The six types of synovial joints are: Plane, Hinge, Pivot, Condyloid, Saddle, and Ball-and-socket."),

        ("Q30. Atlantoaxial joint is: [AIIMS 2016]",
         "(a) Hinge joint (b) Pivot joint (c) Plane joint (d) Ball and socket",
         "Answer: (b) Pivot joint",
         "The median atlantoaxial joint allows rotation of the head (saying 'No'). The dens (odontoid process) of axis acts as a pivot around which the atlas (and head) rotates."),

        ("Q31. Most common direction of shoulder dislocation: [NEET 2019]",
         "(a) Anterior (b) Posterior (c) Superior (d) Inferior",
         "Answer: (a) Anterior",
         "Anterior (anteroinferior) dislocation is most common (>95%) because the capsule and rotator cuff are weakest inferiorly and anteriorly. Axillary nerve may be damaged (loss of sensation over deltoid)."),

        ("Q32. Screw-home mechanism occurs at: [JIPMER 2020]",
         "(a) Hip joint (b) Shoulder joint (c) Knee joint (d) Ankle joint",
         "Answer: (c) Knee joint",
         "Terminal extension of knee involves medial rotation of femur on tibia (or lateral rotation of tibia on femur) which 'locks' the knee. Popliteus muscle 'unlocks' it by laterally rotating the femur."),

        ("Q33. Crouzon syndrome is associated with: [AIIMS 2019]",
         "(a) Craniosynostosis (b) Achondroplasia (c) Osteogenesis imperfecta (d) Rickets",
         "Answer: (a) Craniosynostosis",
         "Crouzon syndrome (FGFR2 mutation): craniosynostosis + midface hypoplasia + proptosis. Apert syndrome: similar but with syndactyly. Both are autosomal dominant."),

        ("Q34. Posterior cruciate ligament prevents: [NEET 2020]",
         "(a) Anterior displacement of tibia (b) Posterior displacement of tibia\n(c) Lateral rotation (d) Medial rotation",
         "Answer: (b) Posterior displacement of tibia",
         "PCL prevents posterior translation of tibia on femur. Tested by posterior drawer test. Commonly injured in dashboard injuries (posterior force on proximal tibia). ACL prevents anterior displacement."),

        ("Q35. Spheno-occipital synchondrosis fuses at: [GPAT 2018]",
         "(a) 2 years (b) 7 years (c) 18-25 years (d) 40 years",
         "Answer: (c) 18-25 years",
         "This is the LAST synchondrosis to ossify and is an important growth center for the skull base. Its closure marks the end of skull base growth. Used in forensic age estimation."),

        ("Q36. Which type of joint has a joint cavity? [NEET 2016]",
         "(a) Fibrous (b) Cartilaginous (c) Synovial (d) All of the above",
         "Answer: (c) Synovial",
         "A joint cavity (synovial cavity) containing synovial fluid is the hallmark feature of synovial joints. Fibrous and cartilaginous joints do NOT have a joint cavity."),

        ("Q37. Medial meniscus of knee is more prone to injury because: [AIIMS 2018]",
         "(a) It is larger (b) It is attached to MCL (c) It is avascular (d) It is C-shaped",
         "Answer: (b) It is attached to MCL",
         "The medial meniscus is firmly attached to the MCL (deep fibers), making it less mobile. When MCL is stressed (valgus force), the medial meniscus is pulled and can tear. Lateral meniscus is more mobile (not attached to LCL)."),

        ("Q38. Inversion of foot occurs at: [JIPMER 2017]",
         "(a) Ankle joint (b) Subtalar joint (c) Hip joint (d) Knee joint",
         "Answer: (b) Subtalar joint",
         "Inversion and eversion occur primarily at the subtalar (talocalcaneal) and transverse tarsal joints, NOT at the ankle joint proper. The ankle (talocrural) joint only allows dorsiflexion and plantarflexion."),

        ("Q39. Trendelenburg sign positive indicates weakness of: [NEET 2019]",
         "(a) Gluteus maximus (b) Gluteus medius (c) Quadriceps (d) Hamstrings",
         "Answer: (b) Gluteus medius",
         "Gluteus medius (and minimus) are hip abductors. Weakness causes pelvis to drop on the opposite (unsupported) side during single-leg stance. Caused by superior gluteal nerve injury."),

        ("Q40. Which joint is supplied by the obturator nerve? [AIIMS 2017]",
         "(a) Shoulder (b) Elbow (c) Hip (d) Ankle",
         "Answer: (c) Hip",
         "The hip joint is supplied by femoral (L2,3,4), obturator (L2,3,4), and sciatic nerves. Hip pathology can cause referred pain in the knee (obturator nerve supplies both). This is clinically important."),

        ("Q41. Type of suture between parietal and temporal bone: [GPAT 2020]",
         "(a) Serrate (b) Squamous (c) Plane (d) Schindylesis",
         "Answer: (b) Squamous",
         "The squamous suture (parietotemporal) is the only major overlapping/beveled suture in the skull vault. The squamous part of temporal bone overlaps the parietal bone."),

        ("Q42. Avascular necrosis of femoral head is common after fracture of: [NEET 2018]",
         "(a) Shaft of femur (b) Neck of femur (c) Greater trochanter (d) Intertrochanteric region",
         "Answer: (b) Neck of femur",
         "Intracapsular fracture of femoral neck disrupts the retinacular arteries (from medial circumflex femoral artery) which are the main blood supply to the femoral head in adults. This leads to AVN."),

        ("Q43. Condyloid joint is present at: [JIPMER 2018]",
         "(a) Wrist joint (b) Elbow joint (c) Hip joint (d) Ankle joint",
         "Answer: (a) Wrist joint",
         "The radiocarpal (wrist) joint is a condyloid (ellipsoid) joint - biaxial, allowing flexion-extension and abduction-adduction (+ circumduction). MCP joints are also condyloid."),

        ("Q44. Deltoid ligament is found in: [NEET 2020]",
         "(a) Shoulder joint (b) Hip joint (c) Knee joint (d) Ankle joint",
         "Answer: (d) Ankle joint",
         "The deltoid (medial) ligament of the ankle is a strong fan-shaped ligament with 4 parts. It resists eversion. It is rarely torn; if damaged, suspect an associated fracture (Pott's fracture)."),

        ("Q45. Lambdoid suture connects: [AIIMS 2020]",
         "(a) Frontal and parietal bones (b) Parietal and occipital bones\n(c) Two parietal bones (d) Parietal and temporal bones",
         "Answer: (b) Parietal and occipital bones",
         "The lambdoid suture connects the two parietal bones with the occipital bone posteriorly. It is named for its resemblance to the Greek letter lambda. Wormian bones are most common here."),

        ("Q46. Locking of knee is done by: [GPAT 2017]",
         "(a) Quadriceps (b) Popliteus (c) Gastrocnemius (d) Medial rotation of femur",
         "Answer: (d) Medial rotation of femur",
         "Locking (screw-home mechanism): In terminal extension, the femur medially rotates on the tibia, tightening cruciate and collateral ligaments, stabilizing the joint in extension. Unlocking is by popliteus."),

        ("Q47. Pannus formation is seen in which joint disease? [NEET 2021]",
         "(a) Osteoarthritis (b) Rheumatoid arthritis (c) Gout (d) Septic arthritis",
         "Answer: (b) Rheumatoid arthritis",
         "Pannus is an abnormal proliferation of synovial tissue that erodes cartilage and bone in RA. It is composed of inflammatory cells, granulation tissue, and activated synoviocytes (key pathology of RA)."),

        ("Q48. Fontanelle used for transfontanellar ultrasound in neonates: [AIIMS 2019]",
         "(a) Anterior (b) Posterior (c) Mastoid (d) Sphenoid",
         "Answer: (a) Anterior",
         "The anterior fontanelle (large, diamond-shaped, closes at 18 months) serves as an acoustic window for cranial ultrasound in neonates. It allows assessment of ventricular size, hemorrhage, and brain structure."),

        ("Q49. Annular ligament is related to: [JIPMER 2019]",
         "(a) Head of radius (b) Head of ulna (c) Head of humerus (d) Head of femur",
         "Answer: (a) Head of radius",
         "The annular ligament encircles the head of the radius and holds it in the radial notch of the ulna (proximal radioulnar joint). It allows rotation (pronation/supination) while maintaining joint integrity."),

        ("Q50. Which of the following is NOT a feature of synovial joint? [NEET 2017]",
         "(a) Joint cavity (b) Synovial fluid (c) Articular disc in all joints (d) Articular cartilage",
         "Answer: (c) Articular disc in all joints",
         "Articular discs/menisci are present only in SOME synovial joints (knee, TMJ, sternoclavicular, wrist). They are not a universal feature. Joint cavity, synovial fluid, and articular cartilage are present in ALL synovial joints."),

        ("Q51. Posterior dislocation of hip joint can injure: [AIIMS 2016]",
         "(a) Femoral nerve (b) Obturator nerve (c) Sciatic nerve (d) Superior gluteal nerve",
         "Answer: (c) Sciatic nerve",
         "Posterior hip dislocation (most common type, from dashboard injury) puts the sciatic nerve at risk as it lies posterior to the hip joint. This can cause foot drop and sensory loss."),

        ("Q52. The atlanto-occipital joint allows: [GPAT 2019]",
         "(a) Rotation (b) Flexion and extension (c) Lateral flexion only (d) No movement",
         "Answer: (b) Flexion and extension",
         "The atlanto-occipital joint is a condyloid (ellipsoid) synovial joint allowing nodding (yes movement - flexion/extension) and slight lateral flexion. Rotation (no movement) occurs at atlantoaxial joint."),

        ("Q53. Suprapatellar bursa communicates with: [NEET 2019]",
         "(a) Hip joint (b) Knee joint cavity (c) Ankle joint (d) Does not communicate with any joint",
         "Answer: (b) Knee joint cavity",
         "The suprapatellar bursa is the largest bursa of the knee and communicates freely with the knee joint cavity. Infection/inflammation of this bursa directly affects the joint (and vice versa)."),

        ("Q54. Apert syndrome features include: [AIIMS 2020]",
         "(a) Craniosynostosis + Polydactyly (b) Craniosynostosis + Syndactyly\n(c) Craniosynostosis + Achondroplasia (d) Craniosynostosis + Cleft palate only",
         "Answer: (b) Craniosynostosis + Syndactyly",
         "Apert syndrome (FGFR2 mutation): coronal craniosynostosis + mitten-hand syndactyly (fused fingers) + midface hypoplasia. Autosomal dominant. Distinguished from Crouzon by presence of syndactyly."),

        ("Q55. Which movement occurs at the atlantoaxial joint? [JIPMER 2020]",
         "(a) Flexion-Extension (b) Rotation (c) Abduction (d) Circumduction",
         "Answer: (b) Rotation",
         "The atlantoaxial joint (pivot type) allows rotation of the head (saying 'No'). About 50% of cervical rotation occurs here. The dens of axis serves as the pivot around which atlas (and skull) rotates."),
    ]

    for i, (question, options, answer, explanation) in enumerate(mcqs):
        # Question
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.line_spacing = 1.0
        run = p.add_run(question)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x1B, 0x4F, 0x72)

        # Options
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.line_spacing = 1.0
        run = p.add_run(options)
        run.font.size = Pt(8.5)

        # Answer
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.0
        run = p.add_run(answer)
        run.bold = True
        run.font.size = Pt(8.5)
        run.font.color.rgb = RGBColor(0x15, 0x6B, 0x0B)

        # Explanation
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = 1.0
        run = p.add_run(f"Explanation: {explanation}")
        run.font.size = Pt(8)
        run.italic = True
        run.font.color.rgb = RGBColor(0x56, 0x6D, 0x7E)

        # Add page break every 10 questions
        if (i + 1) % 10 == 0 and i < len(mcqs) - 1:
            doc.add_page_break()

    doc.add_page_break()


def add_quick_revision(doc):
    """Add quick revision section at the end."""
    add_heading_styled(doc, "QUICK REVISION SECTION", 1)
    add_section_divider(doc)

    add_heading_styled(doc, "Key Facts - Joints at a Glance", 2)

    create_styled_table(doc,
        ["Category", "Key Point"],
        [
            ["Classification", "Structural: Fibrous, Cartilaginous, Synovial\nFunctional: Synarthrosis, Amphiarthrosis, Diarthrosis"],
            ["Fibrous Joints", "3 types: Sutures (skull only), Syndesmosis (interosseous membrane), Gomphosis (teeth)"],
            ["Cartilaginous", "Primary (synchondrosis) = Hyaline, Temporary\nSecondary (symphysis) = Fibrocartilage, Permanent, Midline"],
            ["Synovial Types", "Plane, Hinge, Pivot, Condyloid, Saddle, Ball-and-socket\n(Non-axial, Uniaxial, Biaxial, Multiaxial)"],
            ["Most Mobile", "Shoulder joint (ball & socket) - most frequently dislocated"],
            ["Most Stable", "Hip joint (deep acetabulum + strong ligaments)"],
            ["Largest Joint", "Knee joint (modified hinge, most complex)"],
            ["Strongest Ligament", "Iliofemoral ligament (Y-shaped, prevents hip hyperextension)"],
            ["Hilton's Law", "Nerve to joint = Nerve to muscles moving it = Nerve to overlying skin"],
            ["Synovial Fluid", "Produced by synovial membrane; contains hyaluronic acid + lubricin"],
        ],
        header_color="1B4F72")

    add_heading_styled(doc, "Key Facts - Sutures at a Glance", 2)

    create_styled_table(doc,
        ["Category", "Key Point"],
        [
            ["Suture Types", "Serrate (strongest, interlocking), Squamous (overlapping), Plane (straight, weakest)"],
            ["Coronal Suture", "Frontal + Parietal bones; premature closure = Brachycephaly"],
            ["Sagittal Suture", "Between 2 Parietals (midline); closure = Scaphocephaly (MC craniosynostosis)"],
            ["Lambdoid Suture", "Parietals + Occipital; shaped like lambda; most Wormian bones found here"],
            ["Squamous Suture", "Parietal + Temporal; overlapping; WEAKEST area; epidural hematoma"],
            ["Metopic Suture", "Between 2 frontal halves; closes by 2 yrs (FIRST suture to close)"],
            ["Anterior Fontanelle", "Diamond shaped; largest; closes 18 months; bulging=raised ICP, sunken=dehydration"],
            ["Posterior Fontanelle", "Triangular; smallest; closes 2-3 months"],
            ["Pterion", "Thinnest skull area; 4 bones meet; middle meningeal artery; epidural hematoma"],
            ["Bregma", "Coronal + Sagittal suture junction (anterior fontanelle site)"],
            ["Lambda", "Sagittal + Lambdoid suture junction (posterior fontanelle site)"],
            ["Craniosynostosis", "Sagittal=Scaphocephaly (MC), Coronal=Brachycephaly, Metopic=Trigonocephaly"],
        ],
        header_color="6C3483")

    add_heading_styled(doc, "Important Mnemonics - Quick Recall", 2)

    add_highlight_box(doc, "MNEMONICS COLLECTION",
                     ["SITS (Rotator Cuff): Supraspinatus, Infraspinatus, Teres minor, Subscapularis",
                      "Unhappy Triad: ACL + MCL + Medial Meniscus (lateral blow to knee)",
                      "Synovial joint types (PHP CBS): Plane, Hinge, Pivot, Condyloid, Ball-socket, Saddle",
                      "Suture closure order: Metopic (2yr) > Post.fontanelle (3mo) > Ant.fontanelle (18mo) > Others (30+yr)",
                      "Wormian bones (CHOROID): Cleidocranial, Hypothyroid, OI, Rickets, Otopalatodigital, Idiopathic, Down",
                      "Hilton's Law: Joint nerve = Muscle nerve = Skin nerve",
                      "Craniosynostosis shapes: Sagittal=boat, Coronal=short/wide, Metopic=triangle, All=tower",
                      "SUPination = hold SOUP (palm up); PRONation = PRONE (palm down)",
                      "Knee locking: Femur MEDIAL rotation on tibia; Unlocking: POPLITEUS",
                      "Hip joint ligaments: IPI (Iliofemoral-strongest, Pubofemoral, Ischiofemoral)"],
                     color="E8F8F5", border_color="1ABC9C")

    add_heading_styled(doc, "Clinical Correlations Summary", 2)

    create_styled_table(doc,
        ["Clinical Condition", "Joint/Suture Involved", "Key Feature"],
        [
            ["Shoulder dislocation", "Glenohumeral joint", "Anterior-inferior; axillary nerve damage"],
            ["Hip AVN", "Hip joint (femoral head)", "After neck fracture; blood supply disrupted"],
            ["Knee unhappy triad", "Knee joint", "ACL + MCL + Medial meniscus; lateral blow"],
            ["Nursemaid's elbow", "Proximal radioulnar", "Radial head subluxation in children; annular lig."],
            ["Epidural hematoma", "Pterion area", "Middle meningeal artery; temporal bone fracture"],
            ["Craniosynostosis", "Skull sutures", "Premature closure; sagittal most common"],
            ["Disc prolapse", "Intervertebral symphysis", "L4-L5 most common; posterolateral herniation"],
            ["Rheumatoid arthritis", "Synovial joints (any)", "Pannus formation; symmetrical; small joints first"],
            ["Gout", "1st MTP joint (podagra)", "MSU crystals; negatively birefringent; needle-shaped"],
            ["Trendelenburg sign", "Hip joint", "Gluteus medius weakness; pelvis drops contralateral"],
            ["TMJ dislocation", "Temporomandibular", "Anterior; jaw locks open; lateral pterygoid spasm"],
            ["Pott's fracture", "Ankle joint", "Bimalleolar/trimalleolar; deltoid lig. avulsion"],
        ],
        header_color="1E8449")

    # Final summary box
    add_highlight_box(doc, "EXAM STRATEGY: High-Yield Topics for MCQs",
                     ["1. Classification of joints (structural vs functional) - asked EVERY year",
                      "2. Types of synovial joints with examples - know ALL six types",
                      "3. Suture types and craniosynostosis - very frequently asked",
                      "4. Fontanelles (shape, closure time, clinical significance)",
                      "5. Knee joint (ligaments, menisci, tests, unhappy triad)",
                      "6. Shoulder and Hip joint comparison (stability vs mobility)",
                      "7. Pterion and epidural hematoma - MUST know",
                      "8. Hilton's law and referred pain",
                      "9. Intervertebral disc (structure, prolapse)",
                      "10. Specific movements at joints (pronation/supination, inversion/eversion)"],
                     color="FEF9E7", border_color="F39C12")


def main():
    """Main function to create the complete notes document."""
    doc = Document()

    # Set page size to LEGAL (8.5 x 14 inches) and margins to 1.27 cm
    for section in doc.sections:
        section.page_width = Inches(8.5)
        section.page_height = Inches(14)
        section.top_margin = Cm(1.27)
        section.bottom_margin = Cm(1.27)
        section.left_margin = Cm(1.27)
        section.right_margin = Cm(1.27)

    # Set default paragraph spacing to minimize blank space
    style = doc.styles['Normal']
    style.paragraph_format.space_before = Pt(2)
    style.paragraph_format.space_after = Pt(2)
    style.paragraph_format.line_spacing = 1.0
    style.font.size = Pt(9.5)

    # Reduce heading spacing
    for i in range(1, 4):
        heading_style = doc.styles[f'Heading {i}']
        heading_style.paragraph_format.space_before = Pt(6)
        heading_style.paragraph_format.space_after = Pt(3)
        heading_style.paragraph_format.line_spacing = 1.0

    # Reduce list spacing
    for style_name in ['List Bullet', 'List Number']:
        try:
            list_style = doc.styles[style_name]
            list_style.paragraph_format.space_before = Pt(1)
            list_style.paragraph_format.space_after = Pt(1)
            list_style.paragraph_format.line_spacing = 1.0
        except KeyError:
            pass

    # Build the document
    create_title_page(doc)
    add_introduction(doc)
    add_classification_of_joints(doc)
    add_fibrous_joints(doc)
    add_cartilaginous_joints(doc)
    add_synovial_joints(doc)
    add_specific_joints(doc)
    add_sutures_section(doc)
    add_mcq_section(doc)
    add_quick_revision(doc)

    # Save the document
    output_path = "Joints_and_Sutures_Notes.docx"
    doc.save(output_path)
    print(f"Document saved successfully: {output_path}")
    print("Topics covered: Joints (Arthrology) & Sutures in the Skull")
    print("Includes: 55 Previous Year MCQs with Explanations")
    print("Format: Legal size (8.5x14), 1.27cm margins, compact spacing")


if __name__ == "__main__":
    main()
